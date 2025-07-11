import streamlit as st

# Kullanıcının giriş yapıp yapmadığını kontrol et
if not st.session_state.get('logged_in', False):
    st.warning("Bu sayfayı görüntülemek için giriş yapmanız gerekmektedir.")
    st.switch_page("Home_Page.py") # Giriş sayfasına yönlendir
    st.stop() # Sayfanın geri kalan kodunu çalıştırmayı durdur
import streamlit as st
import pandas as pd
import math
import json
import os
import io # In-memory dosya işlemleri için
from openpyxl.utils import get_column_letter # Excel sütun genişliği için
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill, numbers # Excel hizalama, font, kenarlık, dolgu, SAYI FORMATLARI için
from openpyxl import Workbook # Excel dosyası oluşturmak için
from openpyxl.worksheet.table import Table, TableStyleInfo # Excel tablo için
from docx import Document # Word dosyası oluşturmak için
from docx.shared import Inches, Pt # Word için, Point (yazı boyutu)
from docx.enum.text import WD_ALIGN_PARAGRAPH # Word metin hizalaması için

# --- Constants and Settings ---
NUMBER_OF_SCENARIOS = 9
CURRENCIES = ["TL", "USD", "EUR", "GBP"]
SAVE_FILE_NAME = "finans_inputs.json" # For loading default inputs
CURRENCY_SYMBOLS = {
    "TL": "₺",
    "USD": "$",
    "EUR": "€",
    "GBP": "£"
}

# Financial items for display and export (English names) - Updated order for table
FINANCIAL_ITEMS_EN_DISPLAY = {
    'EBITDA': 'EBITDA',
    'DISCOUNTED CASH FLOW': 'Discounted Cash Flow',
    'CUMULATIVE DISCOUNTED CASH FLOW': 'Cumulative Discounted Cash Flow',
    # NPV, Growth, WACC will be shown separately or in summary
}
# Define the order for detailed report rows
DETAILED_REPORT_ROW_ORDER = [
    'EBITDA',
    'DISCOUNTED CASH FLOW',
    'CUMULATIVE DISCOUNTED CASH FLOW'
]


# --- Helper Functions ---

def format_number_with_currency(number, currency_symbol, is_percentage=False, decimals=2):
    """
    Formats a number to a string with Turkish locale style (dot for thousands, comma for decimals)
    and includes a currency symbol or percentage sign.
    
    This function is primarily for display in Streamlit and Word, where human readability
    with thousand separators is preferred.
    For Excel export, raw numbers are used with Excel's own number formatting.
    """
    if pd.isna(number) or number == '':
        return ''
    
    # Ensure 'number' is truly numeric before attempting operations
    try:
        number = float(number)
    except (ValueError, TypeError):
        return str(number) # Return as string if not a valid number
    
    if is_percentage:
        # For percentages, show as integer percent if it's a whole number, otherwise show with decimals
        if number == int(number):
            return f"{int(number)}%" # e.g., 5%
        else:
            # Format with specified decimals, then replace dot with comma for decimals, add percentage sign
            formatted_str = f"{number:.{decimals}f}".replace('.', ',')
            return f"{formatted_str}%" # e.g., 5,25%
    
    # For currency values: format with 2 decimals, use dot for thousands, comma for decimals
    # f"{number:,.{decimals}f}" uses locale-aware formatting, typically dot for thousands, comma for decimals.
    # We need to manually swap for Turkish/EU style:
    formatted_str = f"{number:,.{decimals}f}"
    
    # Manually swap comma and dot to match TR/EU format (thousands dot, comma decimal)
    # Example: 1,234.56 -> 1.234,56
    parts = formatted_str.split('.')
    integer_part = parts[0].replace(',', '.') # Replace comma (from thousands in US style) with dot
    decimal_part = parts[1] if len(parts) > 1 else "" # Get decimal part if exists

    if decimal_part:
        formatted_output = f"{integer_part},{decimal_part}"
    else:
        formatted_output = integer_part # No decimal part, just integer
    
    return f"{formatted_output} {currency_symbol}"

def calculate_npv(ebitda_base, growth_rate, wacc, projection_years):
    """
    Calculates Net Present Value for a single scenario.
    NPV is then divided by the number of projection years.
    """
    cash_flows = []
    
    # Calculate cash flows for each projection year
    for i in range(projection_years):
        # Convert growth rate to decimal for calculation
        current_ebitda = ebitda_base * ((1 + growth_rate / 100) ** i)
        cash_flows.append(current_ebitda)

    # Calculate Discounted Cash Flow (DCF) for each year
    discounted_cash_flows = []
    for i, cf in enumerate(cash_flows):
        # Discount factor now uses 'i' directly, so first year (i=0) has discount factor of 1
        # This makes the first year's DCF equal to its EBITDA.
        discount_factor = (1 + wacc / 100) ** i
        discounted_cf = cf / discount_factor if discount_factor != 0 else cf # Avoid division by zero
        discounted_cash_flows.append(discounted_cf)

    # Calculate Cumulative Discounted Cash Flow
    cumulative_discounted_cash_flows = [sum(discounted_cash_flows[:i+1]) for i in range(len(discounted_cash_flows))]

    # Net Present Value (NPV) is the sum of all discounted cash flows
    npv = sum(discounted_cash_flows)
    
    # **NEW:** Divide NPV by the number of projection years as requested
    if projection_years > 0:
        npv = npv / projection_years
    else:
        npv = 0 # Handle case where projection_years is zero

    return cash_flows, discounted_cash_flows, cumulative_discounted_cash_flows, npv

def create_excel_report(all_scenario_data, currency_symbol):
    """
    Generates an Excel (.xlsx) report for all scenarios.
    Ensures numbers are written as actual numbers for summation and uses Excel's formatting.
    """
    wb = Workbook()
    
    # Summary Sheet
    ws_summary = wb.active
    ws_summary.title = "Summary"

    # Define styles for summary sheet
    header_font = Font(bold=True, size=11)
    data_font = Font(size=10)
    center_aligned = Alignment(horizontal='center', vertical='center')
    right_aligned = Alignment(horizontal='right', vertical='center')
    thin_border = Border(left=Side(style='thin'), 
                         right=Side(style='thin'), 
                         top=Side(style='thin'), 
                         bottom=Side(style='thin'))
    
    # Excel number formats for percentages:
    # To display % sign BEFORE the number, you would technically use custom format like "%\0"
    # But this often requires writing the raw decimal (e.g., 0.05) and can be tricky with openpyxl.
    # The standard '0%' and '0.00%' typically place the % sign AFTER the number based on locale.
    # Forcing it to the front *and* maintaining numeric type for summation is complex without Excel VBA.
    # We will stick to the standard '0%' and '0.00%' which are commonly understood and retain numeric value.
    excel_percentage_format_no_decimals = '0%' # For 5% (no decimals)
    excel_percentage_format_with_decimals = '0.00%' # For 5.25% (2 decimals)
    
    # Excel number format for currency values (allowing summation) with symbol at the end:
    # Example: #,##0.00 "₺" will show 1.234.567,89 ₺ in Turkish Excel
    excel_currency_numeric_format = f'#,##0.00 "{currency_symbol}"'


    # Summary table headers
    summary_headers = ['Scenario', 'Growth (%)', 'WACC (%)', 'NPV']
    ws_summary.append(summary_headers)
    for col_idx, header_text in enumerate(summary_headers, 1):
        cell = ws_summary.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.alignment = center_aligned
        cell.border = thin_border
    
    # Populate summary table
    for scenario_idx, data in enumerate(all_scenario_data):
        # We are writing raw numbers to Excel for calculations
        # And applying Excel's number format
        row_values = [
            f"Scenario {scenario_idx + 1}",
            data['growth_rate'] / 100, # Convert to decimal for Excel's percentage format
            data['wacc'] / 100,        # Convert to decimal for Excel's percentage format
            data['npv']                # Write raw number
        ]
        ws_summary.append(row_values)
        
        # Apply formatting to summary data rows
        for col_idx in range(1, len(summary_headers) + 1):
            cell = ws_summary.cell(row=ws_summary.max_row, column=col_idx)
            cell.font = data_font
            cell.border = thin_border
            if col_idx == 1: # Scenario name
                cell.alignment = center_aligned
            elif col_idx == 2: # Growth
                cell.alignment = right_aligned
                # Check if original growth rate was whole number
                if data['growth_rate'] == int(data['growth_rate']):
                    cell.number_format = excel_percentage_format_no_decimals
                else:
                    cell.number_format = excel_percentage_format_with_decimals
            elif col_idx == 3: # WACC
                cell.alignment = right_aligned
                # Check if original wacc was whole number
                if data['wacc'] == int(data['wacc']):
                    cell.number_format = excel_percentage_format_no_decimals
                else:
                    cell.number_format = excel_percentage_format_with_decimals
            elif col_idx == 4: # NPV
                cell.alignment = right_aligned
                cell.number_format = excel_currency_numeric_format # Apply custom currency format

    # Adjust column widths for summary sheet
    for col_idx in range(1, ws_summary.max_column + 1):
        max_length = 0
        for row_idx in range(1, ws_summary.max_row + 1):
            cell = ws_summary.cell(row=row_idx, column=col_idx)
            if cell.value is not None:
                # To get the display length for width calculation,
                # use the formatted string (similar to Streamlit/Word display)
                # Need to handle potential IndexError for all_scenario_data if row_idx is 1 (header row)
                if row_idx > 1: # For data rows
                    if col_idx == 2: # Growth
                        val_for_display = all_scenario_data[row_idx-2]['growth_rate'] # Get original value for correct formatting
                        cell_value_str = format_number_with_currency(val_for_display, '', is_percentage=True)
                    elif col_idx == 3: # WACC
                        val_for_display = all_scenario_data[row_idx-2]['wacc'] # Get original value for correct formatting
                        cell_value_str = format_number_with_currency(val_for_display, '', is_percentage=True)
                    elif col_idx == 4: # NPV
                        cell_value_str = format_number_with_currency(cell.value, currency_symbol)
                    else: # Scenario name
                        cell_value_str = str(cell.value)
                else: # For header row
                    cell_value_str = str(cell.value)

                if len(cell_value_str) > max_length:
                    max_length = len(cell_value_str)
        column_letter = get_column_letter(col_idx)
        # Reduce multiplier slightly for tighter fit, but keep padding
        adjusted_width = (max_length + 2) * 1.1 
        if adjusted_width > 0:
            ws_summary.column_dimensions[column_letter].width = adjusted_width


    # Individual Scenario Sheets (Years as columns, Items as rows)
    for scenario_idx, scenario_data in enumerate(all_scenario_data):
        ws = wb.create_sheet(title=f"Scenario {scenario_idx + 1}")
        
        # Scenario specific info above the table
        ws.append([f"Scenario {scenario_idx + 1} Details"])
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=scenario_data['projection_years'] + 1)
        ws['A1'].font = header_font
        ws['A1'].alignment = center_aligned
        ws.append([]) # Blank row

        ws.append(["Starting EBITDA:", scenario_data['ebitda_base']])
        ws.cell(row=ws.max_row, column=2).number_format = excel_currency_numeric_format
        
        ws.append(["Growth Rate (%):", scenario_data['growth_rate']]) # Write raw number (e.g., 10 for 10%)
        # Apply integer percentage format if value is whole number, else decimal percentage
        if scenario_data['growth_rate'] == int(scenario_data['growth_rate']):
            ws.cell(row=ws.max_row, column=2).number_format = excel_percentage_format_no_decimals
            ws.cell(row=ws.max_row, column=2).value = scenario_data['growth_rate'] / 100 # Write 0.10 for 10%
        else:
            ws.cell(row=ws.max_row, column=2).number_format = excel_percentage_format_with_decimals
            ws.cell(row=ws.max_row, column=2).value = scenario_data['growth_rate'] / 100

        ws.append(["WACC (%):", scenario_data['wacc']]) # Write raw number (e.g., 3 for 3%)
        if scenario_data['wacc'] == int(scenario_data['wacc']):
            ws.cell(row=ws.max_row, column=2).number_format = excel_percentage_format_no_decimals
            ws.cell(row=ws.max_row, column=2).value = scenario_data['wacc'] / 100
        else:
            ws.cell(row=ws.max_row, column=2).number_format = excel_percentage_format_with_decimals
            ws.cell(row=ws.max_row, column=2).value = scenario_data['wacc'] / 100
        ws.append([]) # Blank row

        # Detailed table headers (Years)
        years = [scenario_data['start_year'] + i for i in range(scenario_data['projection_years'])]
        headers = ['Financial Items'] + years
        ws.append(headers)
        
        for col_idx, header_text in enumerate(headers, 1):
            cell = ws.cell(row=ws.max_row, column=col_idx)
            cell.font = header_font
            cell.alignment = center_aligned
            cell.border = thin_border

        # Prepare data for new table structure
        # Use the specific order for rows
        detailed_data = {}
        detailed_data['EBITDA'] = scenario_data['cash_flows']
        detailed_data['DISCOUNTED CASH FLOW'] = scenario_data['discounted_cash_flows']
        detailed_data['CUMULATIVE DISCOUNTED CASH FLOW'] = scenario_data['cumulative_discounted_cash_flows']
        
        for item_key in DETAILED_REPORT_ROW_ORDER: # Iterate in predefined order
            row_cells_values = [FINANCIAL_ITEMS_EN_DISPLAY.get(item_key, item_key)] # First cell is item name
            for year_data in detailed_data[item_key]:
                row_cells_values.append(year_data) # Write raw number
            ws.append(row_cells_values)

            # Apply formatting to data rows
            current_row_idx = ws.max_row
            ws.cell(row=current_row_idx, column=1).font = data_font # Financial Item name
            ws.cell(row=current_row_idx, column=1).alignment = right_aligned
            ws.cell(row=current_row_idx, column=1).border = thin_border


            for col_idx in range(2, len(headers) + 1): # Start from year columns
                cell = ws.cell(row=current_row_idx, column=col_idx)
                cell.font = data_font
                cell.alignment = right_aligned # Ensure all numeric cells are right-aligned
                cell.border = thin_border
                cell.number_format = excel_currency_numeric_format # All financial values
                
        ws.append([]) # Blank row
        ws.append(["Net Present Value (NPV):", scenario_data['npv']]) # Write raw NPV
        ws.cell(row=ws.max_row, column=2).number_format = excel_currency_numeric_format
        ws.cell(row=ws.max_row, column=1).font = header_font
        ws.cell(row=ws.max_row, column=2).font = header_font

        # Adjust column widths for detail sheet
        for col_idx in range(1, ws.max_column + 1):
            max_length = 0
            for row_idx in range(1, ws.max_row + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                if cell.value is not None:
                    # Use formatted string for length calculation (as per Streamlit/Word display)
                    if col_idx == 1: # Financial item names column
                         cell_value_str = str(cell.value)
                    else: # Formatted financial values or year headers
                        val = cell.value # Get the raw value
                        # Check if this cell corresponds to Growth Rate or WACC (based on row_idx and header content)
                        if isinstance(val, (int, float)):
                            # These magic numbers (row_idx 3 and 4) come from the fixed layout above the table
                            if row_idx == 3: # Row for Growth Rate
                                val_for_display = scenario_data['growth_rate']
                                cell_value_str = format_number_with_currency(val_for_display, '', is_percentage=True)
                            elif row_idx == 4: # Row for WACC
                                val_for_display = scenario_data['wacc']
                                cell_value_str = format_number_with_currency(val_for_display, '', is_percentage=True)
                            else: # Other financial values
                                cell_value_str = format_number_with_currency(val, currency_symbol)
                        else: # Otherwise, it's likely a year header (string)
                            cell_value_str = str(val)
                        
                    if len(cell_value_str) > max_length:
                        max_length = len(cell_value_str)
            column_letter = get_column_letter(col_idx)
            # Reduce multiplier slightly for tighter fit
            adjusted_width = (max_length + 2) * 1.1
            if adjusted_width > 0:
                ws.column_dimensions[column_letter].width = adjusted_width

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.getvalue()


def create_word_report(all_scenario_data, currency_symbol):
    """
    Generates a Word (.docx) report for all scenarios.
    Uses format_number_with_currency for all numeric displays.
    """
    document = Document()
    document.add_heading('Financial Projections Report', level=1)
    
    for scenario_idx, scenario_data in enumerate(all_scenario_data):
        document.add_heading(f'Scenario {scenario_idx + 1}', level=2)
        document.add_paragraph(f"Starting EBITDA: {format_number_with_currency(scenario_data['ebitda_base'], currency_symbol)}")
        document.add_paragraph(f"Growth Rate: {format_number_with_currency(scenario_data['growth_rate'], '', is_percentage=True, decimals=0 if scenario_data['growth_rate'] == int(scenario_data['growth_rate']) else 2)}")
        document.add_paragraph(f"WACC: {format_number_with_currency(scenario_data['wacc'], '', is_percentage=True, decimals=0 if scenario_data['wacc'] == int(scenario_data['wacc']) else 2)}")
        
        # Add table with years as columns and financial items as rows
        years = [scenario_data['start_year'] + i for i in range(scenario_data['projection_years'])]
        headers = ['Financial Items'] + years
        
        table = document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = str(header_text)
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

        # Prepare data for new table structure
        detailed_data = {}
        detailed_data['EBITDA'] = scenario_data['cash_flows']
        detailed_data['DISCOUNTED CASH FLOW'] = scenario_data['discounted_cash_flows']
        detailed_data['CUMULATIVE DISCOUNTED CASH FLOW'] = scenario_data['cumulative_discounted_cash_flows']

        # Data Rows
        for item_key in DETAILED_REPORT_ROW_ORDER: # Iterate in predefined order
            row_cells = table.add_row().cells
            row_cells[0].text = FINANCIAL_ITEMS_EN_DISPLAY.get(item_key, item_key) # First cell is item name
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            for i, year_data in enumerate(detailed_data[item_key]):
                row_cells[i+1].text = format_number_with_currency(year_data, currency_symbol)
                row_cells[i+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            for cell in row_cells: # Set font size for all cells
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
        
        document.add_paragraph("\n") # Add blank line after table
        document.add_paragraph(f"Net Present Value (NPV): {format_number_with_currency(scenario_data['npv'], currency_symbol)}")
        document.add_paragraph("\n") # Add blank line before next scenario

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()


# --- Streamlit Application Layout ---
st.set_page_config(layout="wide", page_title="Financial Projection Wizard")

st.title("Financial Projection Wizard")
st.markdown("Enter financial parameters and analyze different growth and WACC scenarios.")

# --- Session State Initialization ---
if 'all_scenario_data_for_export' not in st.session_state:
    st.session_state.all_scenario_data_for_export = []
if 'npv_summary_df' not in st.session_state:
    st.session_state.npv_summary_df = pd.DataFrame() # Store DataFrame for summary table

# Load default inputs if available
def load_default_inputs():
    if os.path.exists(SAVE_FILE_NAME):
        try:
            with open(SAVE_FILE_NAME, 'r') as f:
                return json.load(f)
        except json.JSONDecodeError:
            st.error(f"Error reading {SAVE_FILE_NAME}. It might be corrupted. Using default values.")
            return None
    return None

def save_default_inputs(inputs):
    try:
        with open(SAVE_FILE_NAME, 'w') as f:
            json.dump(inputs, f, indent=4)
        st.success("Default inputs saved successfully!")
    except IOError as e:
        st.error(f"Error saving {SAVE_FILE_NAME}: {e}")

default_inputs = load_default_inputs()

# --- Input Section ---
st.header("General Parameters")

col1, col2 = st.columns(2)

with col1:
    start_year_val = int(default_inputs['start_year']) if default_inputs and 'start_year' in default_inputs and str(default_inputs['start_year']).isdigit() else pd.to_datetime('today').year
    start_year = st.number_input("Start Year:", min_value=1900, value=start_year_val, step=1, key="start_year_input", format="%d")

    projection_years_val = int(default_inputs['projection_years']) if default_inputs and 'projection_years' in default_inputs and str(default_inputs['projection_years']).isdigit() else 10
    projection_years = st.number_input("Number of Projection Years:", min_value=1, value=projection_years_val, step=1, key="projection_years_input", format="%d")

with col2:
    selected_currency_val = default_inputs['selected_currency'] if default_inputs and 'selected_currency' in default_inputs else "TL"
    selected_currency = st.selectbox("Currency:", options=CURRENCIES, index=CURRENCIES.index(selected_currency_val) if selected_currency_val in CURRENCIES else 0, key="currency_select")
    
    try:
        ebitda_base_val = float(default_inputs['single_ebitda']) if default_inputs and 'single_ebitda' in default_inputs else 25070432.0
    except (ValueError, TypeError):
        ebitda_base_val = 25070432.0 # Fallback if json value is invalid
    
    # Format for EBITDA input, ensuring it's readable with thousands and decimals
    # Streamlit number_input format="%.2f" allows 2 decimals.
    # For custom thousand separators in input, we'd need a text input and custom parsing.
    # Sticking with default number_input format for now for simpler input.
    ebitda_base = st.number_input("Starting EBITDA (Year 0):", min_value=0.0, value=ebitda_base_val, step=1000.0, format="%.2f", key="ebitda_base_input")


st.header("Scenario Inputs (Growth Rate % and WACC %)")

growth_vars = []
wacc_vars = []

# Dynamically create input fields for 9 scenarios
for i in range(NUMBER_OF_SCENARIOS):
    col = st.columns(1)[0] # Use a single column for each row of inputs
    with col:
        st.markdown(f"**Scenario {i+1}:**")
        
        try:
            default_growth = float(default_inputs['growth_vars'][i]) if default_inputs and 'growth_vars' in default_inputs and i < len(default_inputs['growth_vars']) else (10.0 if i % 3 == 0 else (20.0 if i % 3 == 1 else 30.0))
        except (ValueError, TypeError):
            default_growth = (10.0 if i % 3 == 0 else (20.0 if i % 3 == 1 else 30.0)) # Fallback if json value is invalid
        growth_vars.append(
            st.number_input(f"Growth Rate {i+1} (%):", min_value=-50.0, max_value=100.0, value=default_growth, step=0.1, format="%.1f", key=f"growth_input_{i}")
        )
        
        try:
            default_wacc = float(default_inputs['wacc_vars'][i]) if default_inputs and 'wacc_vars' in default_inputs and i < len(default_inputs['wacc_vars']) else (3.0 if i < 3 else (7.0 if i < 6 else 12.0))
        except (ValueError, TypeError):
            default_wacc = (3.0 if i < 3 else (7.0 if i % 3 == 1 else 12.0)) # Fallback if json value is invalid
        wacc_vars.append(
            st.number_input(f"WACC {i+1} (%):", min_value=0.1, max_value=50.0, value=default_wacc, step=0.1, format="%.1f", key=f"wacc_input_{i}")
        )
        st.markdown("---")

# --- Actions (Calculate, Save/Load Defaults) ---
st.header("Actions")
col_actions1, col_actions2, col_actions3 = st.columns(3)

with col_actions1:
    calculate_button = st.button("Calculate All Scenarios", key="calculate_btn")

with col_actions2:
    if st.button("Save Current Inputs as Default", key="save_defaults_btn"):
        current_inputs = {
            'start_year': str(start_year),
            'projection_years': str(projection_years),
            'selected_currency': selected_currency,
            'single_ebitda': str(ebitda_base),
            'growth_vars': [str(g) for g in growth_vars],
            'wacc_vars': [str(w) for w in wacc_vars]
        }
        save_default_inputs(current_inputs)

with col_actions3:
    if st.button("Load Default Inputs", key="load_defaults_btn"):
        st.experimental_rerun()


# --- Calculation Logic ---
if calculate_button:
    st.session_state.all_scenario_data_for_export = []
    summary_data = []
    
    for i in range(NUMBER_OF_SCENARIOS):
        growth = growth_vars[i]
        wacc = wacc_vars[i]
        
        cash_flows, discounted_cash_flows, cumulative_discounted_cash_flows, npv = calculate_npv(
            ebitda_base, growth, wacc, projection_years
        )
        
        scenario_info = {
            'scenario_name': f"Scenario {i+1}",
            'start_year': start_year,
            'projection_years': projection_years,
            'ebitda_base': ebitda_base,
            'growth_rate': growth,
            'wacc': wacc,
            'cash_flows': cash_flows,
            'discounted_cash_flows': discounted_cash_flows,
            'cumulative_discounted_cash_flows': cumulative_discounted_cash_flows,
            'npv': npv
        }
        st.session_state.all_scenario_data_for_export.append(scenario_info)
        
        summary_data.append({
            'Scenario': f"Scenario {i+1}",
            'Growth (%)': growth,
            'WACC (%)': wacc,
            'NPV': npv
        })
    
    st.session_state.npv_summary_df = pd.DataFrame(summary_data)


# --- Display Results ---
if not st.session_state.npv_summary_df.empty:
    st.header("Summary of Scenarios")
    
    # Format the summary DataFrame for display in Streamlit
    npv_summary_df_display = st.session_state.npv_summary_df.copy()
    # Use 0 decimals if value is integer, else 2 for percentages
    npv_summary_df_display['Growth (%)'] = npv_summary_df_display['Growth (%)'].apply(lambda x: format_number_with_currency(x, '', is_percentage=True, decimals=0 if x == int(x) else 2))
    npv_summary_df_display['WACC (%)'] = npv_summary_df_display['WACC (%)'].apply(lambda x: format_number_with_currency(x, '', is_percentage=True, decimals=0 if x == int(x) else 2))
    npv_summary_df_display['NPV'] = npv_summary_df_display['NPV'].apply(lambda x: format_number_with_currency(x, CURRENCY_SYMBOLS.get(selected_currency, '')))
    
    # Apply styling separately for numeric columns and text columns
    st.dataframe(npv_summary_df_display.style.set_properties(
        subset=['Growth (%)', 'WACC (%)', 'NPV'],
        **{'text-align': 'right'}
    ).set_properties(
        subset=['Scenario'], # Scenario column is left aligned
        **{'text-align': 'left'}
    ), use_container_width=True)

    st.markdown("---")

    st.header("Detailed Scenario Projections")
    tabs = st.tabs([f"Scenario {i+1}" for i in range(NUMBER_OF_SCENARIOS)])
    
    for i, tab in enumerate(tabs):
        with tab:
            if i < len(st.session_state.all_scenario_data_for_export):
                scenario_data = st.session_state.all_scenario_data_for_export[i]
                st.subheader(f"{scenario_data['scenario_name']} Details")
                st.write(f"**Starting EBITDA:** {format_number_with_currency(scenario_data['ebitda_base'], CURRENCY_SYMBOLS.get(selected_currency, ''))}")
                st.write(f"**Growth Rate:** {format_number_with_currency(scenario_data['growth_rate'], '', is_percentage=True, decimals=0 if scenario_data['growth_rate'] == int(scenario_data['growth_rate']) else 2)}")
                st.write(f"**WACC:** {format_number_with_currency(scenario_data['wacc'], '', is_percentage=True, decimals=0 if scenario_data['wacc'] == int(scenario_data['wacc']) else 2)}")
                st.write(f"**Net Present Value (NPV):** {format_number_with_currency(scenario_data['npv'], CURRENCY_SYMBOLS.get(selected_currency, ''))}")
                
                # Prepare data for new table structure for Streamlit display
                years = [scenario_data['start_year'] + yr_idx for yr_idx in range(scenario_data['projection_years'])]
                
                # Create a dictionary where keys are financial items and values are lists of yearly data
                detailed_df_data_raw = {
                    'EBITDA': scenario_data['cash_flows'],
                    'DISCOUNTED CASH FLOW': scenario_data['discounted_cash_flows'],
                    'CUMULATIVE DISCOUNTED CASH FLOW': scenario_data['cumulative_discounted_cash_flows']
                }
                
                # Create a list of dictionaries for DataFrame creation
                df_rows = []
                for item_key in DETAILED_REPORT_ROW_ORDER:
                    row_dict = {'Financial Items': FINANCIAL_ITEMS_EN_DISPLAY.get(item_key, item_key)}
                    for yr_idx, year_val in enumerate(years):
                        row_dict[str(year_val)] = detailed_df_data_raw[item_key][yr_idx]
                    df_rows.append(row_dict)
                
                detailed_df = pd.DataFrame(df_rows)
                
                # Format the values for display in Streamlit DataFrame
                for col_name in detailed_df.columns:
                    if col_name != 'Financial Items':
                        detailed_df[col_name] = detailed_df[col_name].apply(lambda x: format_number_with_currency(x, CURRENCY_SYMBOLS.get(selected_currency, '')))

                # Apply styling separately for numeric columns and text columns
                st.dataframe(detailed_df.style.set_properties(
                    subset=pd.IndexSlice[:, detailed_df.columns[1:]], # All columns except the first one ('Financial Items')
                    **{'text-align': 'right'}
                ).set_properties(
                    subset=pd.IndexSlice[:, detailed_df.columns[0]], # Only the first column ('Financial Items')
                    **{'text-align': 'left'}
                ), use_container_width=True)

    st.markdown("---")

    # --- Download Buttons ---
    st.subheader("Download Results")
    col_dl1, col_dl2 = st.columns(2)

    with col_dl1:
        excel_data = create_excel_report(st.session_state.all_scenario_data_for_export, CURRENCY_SYMBOLS.get(selected_currency, ''))
        st.download_button(
            label="Download All Scenarios as Excel",
            data=excel_data,
            file_name="financial_projections.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="download_excel_btn"
        )

    with col_dl2:
        word_data = create_word_report(st.session_state.all_scenario_data_for_export, CURRENCY_SYMBOLS.get(selected_currency, ''))
        st.download_button(
            label="Download All Scenarios as Word",
            data=word_data,
            file_name="financial_projections.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_word_btn"
        )