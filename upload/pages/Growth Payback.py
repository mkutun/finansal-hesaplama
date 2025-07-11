# dcf_streamlit_app.py - DCF & Credit Analysis Tool with Streamlit
import streamlit as st

# Kullanıcının giriş yapıp yapmadığını kontrol et
if not st.session_state.get('logged_in', False):
    st.warning("Bu sayfayı görüntülemek için giriş yapmanız gerekmektedir.")
    st.switch_page("Home_Page.py") # Giriş sayfasına yönlendir
    st.stop() # Sayfanın geri kalan kodunu çalıştırmayı durdur
import streamlit as st
import pandas as pd
import json
import os
import io # In-memory file operations
from openpyxl.utils import get_column_letter # For Excel column width
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill # For Excel alignment, font, border, fill
from openpyxl import Workbook # For creating Excel files
from openpyxl.worksheet.table import Table, TableStyleInfo # For Excel table
from docx import Document # For creating Word files
from docx.shared import Inches, Pt # For Word, Point (font size)
from docx.enum.text import WD_ALIGN_PARAGRAPH # For Word text alignment

# --- Constants and Settings ---
NUMBER_OF_DISCOUNT_RATES_PER_GROWTH = 3 
SAVE_FILE_NAME = "dcf_streamlit_inputs.json"
CURRENCIES = ["TL", "USD", "EUR", "GBP"]
CURRENCY_SYMBOLS = {
    "TL": "₺",
    "USD": "$",
    "EUR": "€",
    "GBP": "£"
}

# Define the financial items to be displayed and exported, in desired order
FINANCIAL_ITEMS_ORDER = [
    'Discounted Cash Flow (DCF)',
    'DCF Moves Next Year',
    'Total Credit',
    'Subtotal',
    'Working Capital',
    'Amount After Working Capital',
    'Loan Refund Payment',
    'Last Row Total'
]

# --- Excel Report Function (Modified for DCF Data) ---
def create_excel_report(all_scenario_results, currency_symbol):
    try:
        output = io.BytesIO()
        wb = Workbook()

        if 'Sheet' in wb.sheetnames:
            wb.remove(wb['Sheet'])

        border_thin = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

        for scenario_idx, scenario_res in enumerate(all_scenario_results):
            growth_r = scenario_res['growth_rate']
            discount_r = scenario_res['discount_rate']
            data = scenario_res['data']
            
            sheet_name = f"Scenario {scenario_idx+1}"
            ws = wb.create_sheet(sheet_name, index=scenario_idx)

            # Scenario Title
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(data) + 1)
            cell_title = ws.cell(row=1, column=1, value=f"Scenario {scenario_idx+1}: Growth Rate {growth_r:.0f}% / Discount Rate (WACC) {discount_r:.0f}%")
            cell_title.font = Font(bold=True, size=14)
            cell_title.alignment = Alignment(horizontal='center', vertical='center')

            # Input Parameters Section
            ws.cell(row=3, column=1, value="General Inputs:").font = Font(bold=True)
            ws.cell(row=4, column=1, value="Projection Start Year:").font = Font(bold=True)
            ws.cell(row=4, column=2, value=scenario_res['start_year'])
            ws.cell(row=5, column=1, value="Total Projection Years:").font = Font(bold=True)
            ws.cell(row=5, column=2, value=scenario_res['total_simulation_years'])
            ws.cell(row=6, column=1, value="Currency:").font = Font(bold=True)
            ws.cell(row=6, column=2, value=scenario_res['selected_currency'])
            if scenario_res['selected_currency'] != "TL":
                ws.cell(row=7, column=1, value="Exchange Rate:").font = Font(bold=True)
                ws.cell(row=7, column=2, value=f"1 {scenario_res['selected_currency']} = {scenario_res['exchange_rate']:.2f} TL")
            
            ws.cell(row=9, column=1, value="Financial Inputs:").font = Font(bold=True)
            ws.cell(row=10, column=1, value="Initial Discounted Cash Flow (DCF):").font = Font(bold=True)
            ws.cell(row=10, column=2, value=scenario_res['initial_dcf']).number_format = f'{currency_symbol} #,##0.00'
            ws.cell(row=11, column=1, value="Initial Credit Amount:").font = Font(bold=True)
            ws.cell(row=11, column=2, value=scenario_res['initial_credit']).number_format = f'{currency_symbol} #,##0.00'
            ws.cell(row=12, column=1, value="Annual Loan Payment:").font = Font(bold=True)
            ws.cell(row=12, column=2, value=scenario_res['annual_loan_payment']).number_format = f'{currency_symbol} #,##0.00'
            ws.cell(row=13, column=1, value="Loan Term (Years):").font = Font(bold=True)
            ws.cell(row=13, column=2, value=scenario_res['loan_term_years'])
            ws.cell(row=14, column=1, value="Grace Period (Years):").font = Font(bold=True)
            ws.cell(row=14, column=2, value=scenario_res['grace_period_years'])

            current_row = 16 # Start of the main data table

            years_list = sorted(list(data.keys()))
            
            # Header Row
            ws.cell(row=current_row, column=1, value="Financial Item").font = Font(bold=True, color="FFFFFF")
            ws.cell(row=current_row, column=1).fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
            ws.cell(row=current_row, column=1).border = border_thin

            for col_idx, year_val in enumerate(years_list, 2):
                cell = ws.cell(row=current_row, column=col_idx, value=year_val)
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.border = border_thin
            
            # Data Rows
            for row_num, item_name in enumerate(FINANCIAL_ITEMS_ORDER, current_row + 1):
                ws.cell(row=row_num, column=1, value=item_name).font = Font(bold=True)
                ws.cell(row=row_num, column=1).alignment = Alignment(horizontal='left', vertical='center')
                ws.cell(row=row_num, column=1).border = border_thin

                for col_idx, year_val in enumerate(years_list, 2): # Start from column 2
                    value = data[year_val].get(item_name, 0)
                    cell = ws.cell(row=row_num, column=col_idx, value=value)
                    cell.alignment = Alignment(horizontal='right', vertical='center')
                    cell.border = border_thin

                    # Apply number formatting
                    if item_name in ['Discounted Cash Flow (DCF)', 'DCF Moves Next Year', 'Total Credit',
                                    'Subtotal', 'Amount After Working Capital', 'Loan Refund Payment', 'Last Row Total']:
                        cell.number_format = f'{currency_symbol} #,##0.00'
                    else:
                        cell.number_format = '#,##0.00' # Default for others

            # Adjust column widths
            for col_idx in range(1, len(years_list) + 2): # +2 for "Financial Item" column and 1-based indexing
                max_length = 0
                for r_idx in range(1, ws.max_row + 1):
                    cell_value = ws.cell(row=r_idx, column=col_idx).value
                    if cell_value is not None:
                        # Account for currency symbols and formatting in length calculation
                        if isinstance(cell_value, (int, float)):
                            if ws.cell(row=r_idx, column=col_idx).number_format == f'{currency_symbol} #,##0.00':
                                formatted_str = f"{cell_value:,.2f} {currency_symbol}"
                            else:
                                formatted_str = f"{cell_value:,.2f}"
                        else:
                            formatted_str = str(cell_value)
                        
                        max_length = max(max_length, len(formatted_str))
                adjusted_width = (max_length + 2) * 1.2 # Add padding
                ws.column_dimensions[get_column_letter(col_idx)].width = adjusted_width

        wb.save(output)
        output.seek(0)
        return output.getvalue()
    except Exception as e:
        st.error(f"An error occurred while creating the Excel report: {e}")
        st.exception(e) # Show full traceback in Streamlit debug area
        return None # Return empty bytes on error

# --- Word Report Function (Modified for DCF Data) ---
def create_word_report(all_scenario_results, currency_symbol):
    try:
        output = io.BytesIO()
        document = Document()

        document.add_heading("Financial Projection and Credit Analysis Report", level=1)
        document.add_paragraph(f"Report Date: {pd.to_datetime('today').strftime('%Y-%m-%d %H:%M')}")
        document.add_paragraph("\n")

        for scenario_idx, scenario_res in enumerate(all_scenario_results):
            growth_r = scenario_res['growth_rate']
            discount_r = scenario_res['discount_rate']
            data = scenario_res['data']

            document.add_heading(f"Scenario {scenario_idx+1}: Growth Rate {growth_r:.0f}% / Discount Rate (WACC) {discount_r:.0f}%", level=2)
            
            # Add Input Parameters to Word
            document.add_paragraph("General Inputs:")
            document.add_paragraph(f"• Projection Start Year: {scenario_res['start_year']}")
            document.add_paragraph(f"• Total Projection Years: {scenario_res['total_simulation_years']}")
            document.add_paragraph(f"• Currency: {scenario_res['selected_currency']}")
            if scenario_res['selected_currency'] != "TL":
                document.add_paragraph(f"• Exchange Rate: 1 {scenario_res['selected_currency']} = {scenario_res['exchange_rate']:.2f} TL")
            
            document.add_paragraph("\nFinancial Inputs:")
            document.add_paragraph(f"• Initial Discounted Cash Flow (DCF): {scenario_res['initial_dcf']:,.2f} {currency_symbol}")
            document.add_paragraph(f"• Initial Credit Amount: {scenario_res['initial_credit']:,.2f} {currency_symbol}")
            document.add_paragraph(f"• Annual Loan Payment: {scenario_res['annual_loan_payment']:,.2f} {currency_symbol}")
            document.add_paragraph(f"• Loan Term (Years): {scenario_res['loan_term_years']}")
            document.add_paragraph(f"• Grace Period (Years): {scenario_res['grace_period_years']}")
            document.add_paragraph("\n")

            years_list = sorted(list(data.keys()))
            num_rows = len(FINANCIAL_ITEMS_ORDER) + 1 # Items + Header
            num_cols = len(years_list) + 1 # Years + Item column

            table = document.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'

            # Header Row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Financial Item"
            hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            hdr_cells[0].paragraphs[0].runs[0].font.bold = True
            hdr_cells[0].paragraphs[0].runs[0].font.size = Pt(10)

            for j, year in enumerate(years_list):
                cell = hdr_cells[j + 1]
                cell.text = str(year)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(10)

            # Data Rows
            for i, item_name in enumerate(FINANCIAL_ITEMS_ORDER):
                row_cells = table.rows[i + 1].cells
                row_cells[0].text = item_name
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row_cells[0].paragraphs[0].runs[0].font.bold = True
                row_cells[0].paragraphs[0].runs[0].font.size = Pt(9)

                for j, year in enumerate(years_list):
                    cell = row_cells[j + 1]
                    value = data[year].get(item_name, 0)
                    display_val = ""
                    
                    if pd.isna(value) or value == '':
                        display_val = ''
                    elif item_name in ['Discounted Cash Flow (DCF)', 'DCF Moves Next Year', 'Total Credit',
                                    'Subtotal', 'Amount After Working Capital', 'Loan Refund Payment', 'Last Row Total']:
                        display_val = f"{value:,.2f} {currency_symbol}"
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else: # For Working Capital or other numeric values if they ever change from 0
                        display_val = f"{value:,.2f}"
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                    cell.text = display_val
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
            
            document.add_page_break() # New page for each scenario

        document.save(output)
        output.seek(0)
        return output.getvalue()
    except Exception as e:
        st.error(f"An error occurred while creating the Word report: {e}")
        st.exception(e) # Show full traceback in Streamlit debug area
        return None # Return empty bytes on error


# --- Calculation Logic (Murat's Confirmed Excel Formulas) ---
def calculate_dcf_and_credit(growth_rate, discount_rate, initial_dcf,
                             initial_credit, annual_loan_payment, loan_term_years, grace_period_years,
                             start_year, total_simulation_years):
    """
    Performs DCF and Credit calculations based on provided parameters.
    Returns a dictionary containing calculated data for each year.
    All formulas are now aligned with Murat's Excel logic.
    """
    results = {}
    current_dcf_value = initial_dcf 
    current_credit_balance = initial_credit # Credit is entered as negative (debt)

    prev_dcf_moves_next_year_value = 0 # Initial value for the first year's calculation of DCF Moves Next Year

    for i in range(total_simulation_years):
        year = start_year + i
        year_data = {}

        # 1. Discounted Cash Flow (DCF) Calculation
        if i == 0: # First year
            current_dcf_value_this_year = initial_dcf
        else: # Subsequent years
            current_dcf_value_this_year = current_dcf_value * (1 + growth_rate) / (1 + discount_rate)
        
        year_data['Discounted Cash Flow (DCF)'] = current_dcf_value_this_year
        current_dcf_value = current_dcf_value_this_year # Update for next iteration's DCF calculation


        # 2. DCF Moves Next Year Calculation
        # Formula: Previous Year's DCF Moves Next Year (or Subtotal) + Current Year's Discounted Cash Flow (DCF)
        year_data['DCF Moves Next Year'] = prev_dcf_moves_next_year_value + year_data['Discounted Cash Flow (DCF)']
        

        # 3. Loan Refund Payment Calculation
        loan_refund_payment_this_year = 0
        if i >= grace_period_years and current_credit_balance < -0.01:
            loan_refund_payment_this_year = min(annual_loan_payment, abs(current_credit_balance))
            
            if i >= (grace_period_years + loan_term_years):
                if abs(current_credit_balance) <= 0.01:
                    loan_refund_payment_this_year = 0
                else:
                    loan_refund_payment_this_year = abs(current_credit_balance)
                
        year_data['Loan Refund Payment'] = loan_refund_payment_this_year

        # 4. Total Credit Calculation
        if i == 0:
            year_data['Total Credit'] = initial_credit
        else:
            current_credit_balance += year_data['Loan Refund Payment']
            
            if current_credit_balance > 0 and year_data['Loan Refund Payment'] > 0:
                current_credit_balance = 0 

            year_data['Total Credit'] = current_credit_balance
            
            if abs(year_data['Total Credit']) < 0.01 and year_data['Loan Refund Payment'] > 0:
                year_data['Total Credit'] = 0


        # 5. Subtotal Calculation (CONFIRMED: Subtotal = DCF Moves Next Year)
        year_data['Subtotal'] = year_data['DCF Moves Next Year']
        
        prev_dcf_moves_next_year_value = year_data['DCF Moves Next Year']


        # 6. Working Capital (Placeholder)
        year_data['Working Capital'] = 0

        # 7. Amount After Working Capital
        year_data['Amount After Working Capital'] = year_data['Subtotal'] - year_data['Working Capital']

        # 8. Last Row (Amount After Working Capital - Loan Refund Payment)
        year_data['Last Row Total'] = year_data['Amount After Working Capital'] - year_data['Loan Refund Payment']

        results[year] = year_data

    return results

# --- Session State Initialization ---
if 'show_results' not in st.session_state:
    st.session_state.show_results = False
if 'all_scenario_results' not in st.session_state:
    st.session_state.all_scenario_results = []
if 'last_inputs' not in st.session_state:
    st.session_state.last_inputs = {}

# --- Load Default Inputs ---
def load_default_inputs():
    if os.path.exists(SAVE_FILE_NAME):
        try:
            with open(SAVE_FILE_NAME, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            st.warning(f"Could not load default inputs: {e}")
    return {}

# --- Save Inputs ---
def save_inputs(inputs):
    try:
        with open(SAVE_FILE_NAME, 'w', encoding='utf-8') as f:
            json.dump(inputs, f, indent=4)
    except Exception as e:
        st.error(f"Could not save inputs: {e}")

# --- Streamlit Application Layout ---
st.set_page_config(layout="wide", page_title="Financial Projection & Credit Analysis")

st.title("Financial Projection & Credit Analysis")
st.markdown("Enter parameters and view financial projections and cash flows under different scenarios.")

default_inputs = load_default_inputs()

# --- Input Section ---
st.header("Input Parameters")

col_params1, col_params2, col_params3 = st.columns(3)

with col_params1:
    st.subheader("General Settings")
    start_year = st.number_input(
        "Projection Start Year:",
        min_value=1900,
        max_value=2100,
        value=int(default_inputs.get("start_year", 2026)),
        step=1,
        key="start_year_input"
    )
    total_simulation_years = st.number_input(
        "Total Number of Projection Years:",
        min_value=1,
        max_value=30,
        value=int(default_inputs.get("total_simulation_years", 10)),
        step=1,
        key="total_simulation_years_input"
    )
    selected_currency = st.selectbox(
        "Select Currency:",
        CURRENCIES,
        index=CURRENCIES.index(default_inputs.get("selected_currency", "TL")) if default_inputs.get("selected_currency", "TL") in CURRENCIES else 0,
        key="currency_select"
    )
    exchange_rate_value = float(default_inputs.get("exchange_rate", 1.0))
    if selected_currency == "TL":
        exchange_rate_value = 1.0

    exchange_rate = st.number_input(
        f"Exchange Rate (for {selected_currency}, 1 {selected_currency} = X TL):",
        min_value=0.01,
        value=exchange_rate_value,
        step=0.01,
        format="%.2f",
        disabled=(selected_currency == "TL"),
        key="exchange_rate_input"
    )

with col_params2:
    st.subheader("Initial Values")
    initial_dcf = st.number_input(
        f"Initial Discounted Cash Flow (DCF) ({CURRENCY_SYMBOLS.get(selected_currency, '')}):",
        min_value=0.0,
        value=float(default_inputs.get("initial_dcf", 367200000)),
        step=100000.0,
        format="%.2f",
        key="initial_dcf_input"
    )
    initial_credit = st.number_input(
        f"Initial Credit Amount (Negative for Debt) ({CURRENCY_SYMBOLS.get(selected_currency, '')}):",
        value=float(default_inputs.get("initial_credit", -3000000000)),
        step=100000.0,
        format="%.2f",
        key="initial_credit_input"
    )

with col_params3:
    st.subheader("Loan Repayment")
    annual_loan_payment = st.number_input(
        f"Annual Loan Payment Amount ({CURRENCY_SYMBOLS.get(selected_currency, '')}):",
        min_value=0.0,
        value=float(default_inputs.get("annual_loan_payment", 500000000)),
        step=100000.0,
        format="%.2f",
        key="annual_loan_payment_input"
    )
    loan_term_years = st.number_input(
        "Loan Term (Years):",
        min_value=0,
        max_value=30,
        value=int(default_inputs.get("loan_term_years", 10)),
        step=1,
        key="loan_term_years_input"
    )
    grace_period_years = st.number_input(
        "Grace Period (Years):",
        min_value=0,
        max_value=loan_term_years, # Grace period cannot exceed loan term
        value=int(default_inputs.get("grace_period_years", 2)),
        step=1,
        key="grace_period_years_input"
    )

st.markdown("---")
st.header("Scenario Definitions")

num_growth_groups = st.number_input("How many different growth rate scenarios?", min_value=1, value=int(default_inputs.get("num_growth_groups", 3)), step=1)
growth_rates_main = []
# Ensure initial load of main_growth_rates_values
if "main_growth_rates_values" not in st.session_state or len(st.session_state.main_growth_rates_values) != num_growth_groups:
    st.session_state.main_growth_rates_values = default_inputs.get("main_growth_rates_values", [15.0, 20.0, 25.0])
    if len(st.session_state.main_growth_rates_values) < num_growth_groups:
        st.session_state.main_growth_rates_values.extend([15.0] * (num_growth_groups - len(st.session_state.main_growth_rates_values)))
    elif len(st.session_state.main_growth_rates_values) > num_growth_groups:
        st.session_state.main_growth_rates_values = st.session_state.main_growth_rates_values[:num_growth_groups]


for i in range(num_growth_groups):
    growth_rates_main.append(
        st.number_input(
            f"Growth Rate Scenario Group {i+1} (%):",
            min_value=-100.0,
            max_value=1000.0,
            value=st.session_state.main_growth_rates_values[i],
            step=1.0,
            format="%.0f",
            key=f"growth_rate_main_{i}"
        )
    )
    st.session_state.main_growth_rates_values[i] = growth_rates_main[i] # Update session state

st.markdown("---")
st.subheader("Discount Rates (WACC) for Each Growth Rate")

# Initialize discount_rates_per_growth if not in session state or if growth rates changed
if "discount_rates_per_growth" not in st.session_state:
    st.session_state.discount_rates_per_growth = default_inputs.get("discount_rates_per_growth", {})

# Clean up stale entries and add new ones based on current growth_rates_main
current_growth_rate_keys = {str(gr) for gr in growth_rates_main}
stale_keys = [k for k in st.session_state.discount_rates_per_growth if k not in current_growth_rate_keys]
for key in stale_keys:
    del st.session_state.discount_rates_per_growth[key]

for i, growth_rate_val in enumerate(growth_rates_main):
    growth_rate_str = str(growth_rate_val)
    if growth_rate_str not in st.session_state.discount_rates_per_growth:
        st.session_state.discount_rates_per_growth[growth_rate_str] = [3.0, 7.0, 12.0][:NUMBER_OF_DISCOUNT_RATES_PER_GROWTH] # Default initial WACCs

    st.write(f"**Discount Rates (WACC) for Growth Rate: {growth_rate_val:.0f}%:**")
    
    num_waccs_for_this_growth = st.number_input(
        f"How many WACC values for Growth {growth_rate_val:.0f}%?",
        min_value=1,
        max_value=5, # Limit for reasonable input
        value=len(st.session_state.discount_rates_per_growth[growth_rate_str]),
        step=1,
        key=f"num_waccs_{i}"
    )
    
    current_wacc_list = st.session_state.discount_rates_per_growth[growth_rate_str]
    if len(current_wacc_list) < num_waccs_for_this_growth:
        current_wacc_list.extend([7.0] * (num_waccs_for_this_growth - len(current_wacc_list)))
    elif len(current_wacc_list) > num_waccs_for_this_growth:
        current_wacc_list = current_wacc_list[:num_waccs_for_this_growth]
    st.session_state.discount_rates_per_growth[growth_rate_str] = current_wacc_list


    cols = st.columns(num_waccs_for_this_growth)
    temp_wacc_inputs = []
    for j in range(num_waccs_for_this_growth):
        with cols[j]:
            wacc_val = st.number_input(
                f"WACC {j+1} (%):",
                min_value=0.0,
                max_value=100.0,
                value=st.session_state.discount_rates_per_growth[growth_rate_str][j],
                step=0.1,
                format="%.0f",
                key=f"wacc_input_growth_{i}_wacc_{j}"
            )
            temp_wacc_inputs.append(wacc_val)
    
    st.session_state.discount_rates_per_growth[growth_rate_str] = temp_wacc_inputs # Update for this growth rate

st.markdown("---")

# --- Calculate Buttons and Display Results ---
col_calc, col_clear = st.columns([1, 0.2])

with col_calc:
    if st.button("Calculate and Show Results", key="calculate_btn"):
        # Save current inputs to JSON for next session
        current_inputs_to_save = {
            "start_year": start_year,
            "total_simulation_years": total_simulation_years,
            "selected_currency": selected_currency,
            "exchange_rate": exchange_rate,
            "initial_dcf": initial_dcf,
            "initial_credit": initial_credit,
            "annual_loan_payment": annual_loan_payment,
            "loan_term_years": loan_term_years,
            "grace_period_years": grace_period_years,
            "num_growth_groups": num_growth_groups,
            "main_growth_rates_values": st.session_state.main_growth_rates_values,
            "discount_rates_per_growth": st.session_state.discount_rates_per_growth # Save the dictionary
        }
        save_inputs(current_inputs_to_save)

        st.session_state.all_scenario_results = []
        
        # Iterate through each growth rate group
        for growth_rate_val in growth_rates_main:
            wacc_list_for_this_growth = st.session_state.discount_rates_per_growth.get(str(growth_rate_val), [])
            
            for discount_rate_val in wacc_list_for_this_growth:
                # Adjust initial_dcf and initial_credit based on exchange rate for calculation
                # Calculations are done in TL base if exchange rate is applied
                calculated_initial_dcf = initial_dcf
                calculated_initial_credit = initial_credit
                calculated_annual_loan_payment = annual_loan_payment

                if selected_currency != "TL" and exchange_rate > 0:
                    calculated_initial_dcf *= exchange_rate
                    calculated_initial_credit *= exchange_rate
                    calculated_annual_loan_payment *= exchange_rate

                scenario_data = calculate_dcf_and_credit(
                    growth_rate=growth_rate_val / 100.0,       # Convert to decimal
                    discount_rate=discount_rate_val / 100.0,   # Convert to decimal
                    initial_dcf=calculated_initial_dcf,
                    initial_credit=calculated_initial_credit,
                    annual_loan_payment=calculated_annual_loan_payment,
                    loan_term_years=loan_term_years,
                    grace_period_years=grace_period_years,
                    start_year=start_year,
                    total_simulation_years=total_simulation_years
                )
                
                st.session_state.all_scenario_results.append({
                    "growth_rate": growth_rate_val,
                    "discount_rate": discount_rate_val,
                    "data": scenario_data,
                    "initial_dcf": initial_dcf, # Keep original input for display/export
                    "initial_credit": initial_credit, # Keep original input for display/export
                    "annual_loan_payment": annual_loan_payment, # Keep original input for display/export
                    "loan_term_years": loan_term_years,
                    "grace_period_years": grace_period_years,
                    "start_year": start_year,
                    "total_simulation_years": total_simulation_years,
                    "selected_currency": selected_currency,
                    "exchange_rate": exchange_rate
                })
        
        st.session_state.show_results = True

with col_clear:
    if st.button("Clear Results", key="clear_btn"):
        st.session_state.show_results = False
        st.session_state.all_scenario_results = []
        st.rerun()

# --- Display Results (controlled by session_state) ---
if st.session_state.show_results:
    st.header("Projection Results")
    currency_symbol = CURRENCY_SYMBOLS.get(selected_currency, '')

    for i, scenario_res in enumerate(st.session_state.all_scenario_results):
        growth_r = scenario_res['growth_rate']
        discount_r = scenario_res['discount_rate']
        data = scenario_res['data']
        
        st.subheader(f"Scenario {i+1}: Growth Rate {growth_r:.0f}% / Discount Rate (WACC) {discount_r:.0f}%")

        # Prepare data for display in DataFrame
        years = sorted(list(data.keys()))
        
        display_data_dict = {}
        for item in FINANCIAL_ITEMS_ORDER:
            row_values = []
            for year in years:
                value = data[year].get(item, 0)
                # Apply currency formatting and symbol for relevant fields
                if item in ['Discounted Cash Flow (DCF)', 'DCF Moves Next Year', 'Total Credit',
                            'Subtotal', 'Amount After Working Capital', 'Loan Refund Payment', 'Last Row Total']:
                    row_values.append(f"{value:,.2f} {currency_symbol}")
                else:
                    row_values.append(f"{value:,.2f}") # For Working Capital (if not 0), or other numeric
            display_data_dict[item] = row_values

        pivoted_df = pd.DataFrame(display_data_dict).T
        pivoted_df.columns = [str(year) for year in years]
        pivoted_df.index.name = 'Financial Item'
        
        st.dataframe(pivoted_df.style.set_properties(**{'text-align': 'right'}), use_container_width=True)
        st.markdown("---")

    st.subheader("Total Values for All Scenarios")
    summary_data = []
    for scenario_idx, scenario_res in enumerate(st.session_state.all_scenario_results):
        growth_r = scenario_res['growth_rate']
        discount_r = scenario_res['discount_rate']
        data = scenario_res['data']
        
        last_year = max(data.keys())
        last_row_total = data[last_year]['Last Row Total']
        
        summary_data.append({
            "Scenario No": f"Scenario {scenario_idx+1}",
            "Growth Rate (%)": f"{growth_r:.0f}%",
            "Discount Rate (WACC) (%)": f"{discount_r:.0f}%",
            f"Last Year ({last_year}) Total ({currency_symbol})": f"{last_row_total:,.2f}"
        })
    
    if summary_data:
        summary_df = pd.DataFrame(summary_data)
        st.dataframe(summary_df.style.set_properties(**{'text-align': 'right'}), use_container_width=True)

    st.markdown("---")
    st.subheader("Download Reports")

    col_dl1, col_dl2 = st.columns(2)

    with col_dl1:
        # Generate Excel data only if results exist
        if st.session_state.all_scenario_results:
            excel_data = create_excel_report(st.session_state.all_scenario_results, CURRENCY_SYMBOLS.get(selected_currency, ''))
            if excel_data: # Only show download button if data was successfully generated
                st.download_button(
                    label="Download All Scenarios as Excel",
                    data=excel_data,
                    file_name="financial_projections.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="download_excel_btn"
                )
            else:
                st.warning("Excel report could not be generated.")

    with col_dl2:
        # Generate Word data only if results exist
        if st.session_state.all_scenario_results:
            word_data = create_word_report(st.session_state.all_scenario_results, CURRENCY_SYMBOLS.get(selected_currency, ''))
            if word_data: # Only show download button if data was successfully generated
                st.download_button(
                    label="Download All Scenarios as Word",
                    data=word_data,
                    file_name="financial_projections.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_word_btn"
                )
            else:
                st.warning("Word report could not be generated.")