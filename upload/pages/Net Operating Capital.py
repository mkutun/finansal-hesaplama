import streamlit as st
import pandas as pd
from io import BytesIO
from openpyxl.workbook import Workbook
from openpyxl.styles import Alignment, Font
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# Sayfa Yapılandırması
st.set_page_config(page_title="Net Working Capital Analysis", layout="centered")

st.title("📊 Net Working Capital Analysis Application")
st.write("Murat, this application will help you calculate your business's net working capital requirement, just as Fercan Hoca taught.")

# --- SESSION STATE İLK DEĞERLERİ BAŞLATMA ---
# Bu kısım, uygulamanın ilk açılışında veya sıfırlandığında çalışır
if 'calculation_successful' not in st.session_state:
    st.session_state.calculation_successful = False
    st.session_state.calculated_data = {} # Hesaplanan tüm verileri burada saklayacağız

# --- CURRENCY SELECTION AND EXCHANGE RATE INPUT ---
st.header("💱 Currency Information")

currency_options = {
    "TL": "₺",
    "USD": "$",
    "EUR": "€",
    "GBP": "£"
}
# Session state kullanarak selectbox'ın değerini koruyalım
selected_currency_name = st.selectbox(
    "Select Currency",
    list(currency_options.keys()),
    key="currency_select",
    help="Choose the currency for your calculations."
)
currency_symbol = currency_options[selected_currency_name]

# Kur giriş alanı ve hesaplama mantığı
exchange_rate_input = 1.0 # Default value for TL

if selected_currency_name != "TL":
    st.info(f"Please enter the exchange rate for 1 {selected_currency_name} to TL. If you enter 1, the calculation will use the input values directly without conversion.")
    # Session state kullanarak number_input'ın değerini koruyalım
    exchange_rate_input = st.number_input(f"Exchange Rate (1 {selected_currency_name} = ? TL)", min_value=0.0, value=st.session_state.get('exchange_rate_input_value', 1.0), format="%.4f", key="exchange_rate_input", help="Enter the current exchange rate for the selected currency against Turkish Lira.")
    
    if exchange_rate_input == 0:
        st.error("Exchange rate cannot be zero. Please enter a valid rate.")
        st.stop()

# Kurun hesaplamaya etkisini belirleme
if selected_currency_name == "TL" or exchange_rate_input == 1.0:
    effective_exchange_rate = 1.0
else:
    effective_exchange_rate = exchange_rate_input


# --- INPUT INFORMATION ---
st.header("📈 Input Information")

# Kullanıcıdan giriş alırken örnek değerleri de verelim
# Bu değerler, resimdeki örnek verilerdir.
default_sales = 70000000.0
default_smm = 35000000.0
default_trade_receivables = 20000000.0
default_inventories = 9000000.0
default_trade_payables = 15000000.0

# Yeni eklenecek manuel girişler
default_current_assets = 30000000.0
default_current_liabilities = 25000000.0


col1, col2 = st.columns(2)

with col1:
    sales_input = st.number_input(f"Annual Sales ({currency_symbol})", min_value=0.0, value=st.session_state.get('sales_input_value', default_sales), step=100000.0, format="%.2f", key="sales_input", help="Total sales of your business for one year in the selected currency.")
    smm_input = st.number_input(f"Cost of Goods Sold (COGS) ({currency_symbol})", min_value=0.0, value=st.session_state.get('smm_input_value', default_smm), step=100000.0, format="%.2f", key="smm_input", help="Cost of goods sold by your business.")
    trade_receivables_input = st.number_input(f"Average Trade Receivables ({currency_symbol})", min_value=0.0, value=st.session_state.get('trade_receivables_input_value', default_trade_receivables), step=10000.0, format="%.2f", key="trade_receivables_input", help="Average amount to be collected from customers.")
    
    current_assets_input = st.number_input(f"Current Assets ({currency_symbol})", min_value=0.0, value=st.session_state.get('current_assets_input_value', default_current_assets), step=10000.0, format="%.2f", key="current_assets_input", help="Total value of current assets (e.g., cash, accounts receivable, inventory).")


with col2:
    inventories_input = st.number_input(f"Average Inventories ({currency_symbol})", min_value=0.0, value=st.session_state.get('inventories_input_value', default_inventories), step=10000.0, format="%.2f", key="inventories_input", help="Average value of inventory held by your business.")
    trade_payables_input = st.number_input(f"Average Trade Payables ({currency_symbol})", min_value=0.0, value=st.session_state.get('trade_payables_input_value', default_trade_payables), step=10000.0, format="%.2f", key="trade_payables_input", help="Average amount owed to suppliers.")
    
    current_liabilities_input = st.number_input(f"Current Liabilities ({currency_symbol})", min_value=0.0, value=st.session_state.get('current_liabilities_input_value', default_current_liabilities), step=10000.0, format="%.2f", key="current_liabilities_input", help="Total value of current liabilities (e.g., accounts payable, short-term debt).")


# Hesapla düğmesi
calculate_button = st.button("Calculate Net Working Capital")

# Sıfırlama butonu (isteğe bağlı ama pratik)
reset_button = st.button("Reset Inputs and Calculations")
if reset_button:
    # st.session_state'i sıfırla
    for key in ['calculation_successful', 'calculated_data', 'exchange_rate_input_value',
                'sales_input_value', 'smm_input_value', 'trade_receivables_input_value',
                'inventories_input_value', 'trade_payables_input_value',
                'current_assets_input_value', 'current_liabilities_input_value']: # Yeni eklenenleri de sıfırla
        if key in st.session_state:
            del st.session_state[key]
    st.rerun() # Uygulamayı yeniden başlat

# Sonuçları göstermek için bir yer tutucu
results_placeholder = st.empty()
download_placeholder = st.empty()


# Hesaplama mantığı
if calculate_button or st.session_state.calculation_successful: # Düğmeye basıldığında veya daha önce hesaplandıysa göster
    # Yeni bir hesaplama yapılacaksa veya ilk defa hesaplanıyorsa
    if calculate_button:
        # Giriş değerlerini session state'e kaydet
        st.session_state.exchange_rate_input_value = exchange_rate_input
        st.session_state.sales_input_value = sales_input
        st.session_state.smm_input_value = smm_input
        st.session_state.trade_receivables_input_value = trade_receivables_input
        st.session_state.inventories_input_value = inventories_input
        st.session_state.trade_payables_input_value = trade_payables_input
        st.session_state.current_assets_input_value = current_assets_input # Yeni
        st.session_state.current_liabilities_input_value = current_liabilities_input # Yeni
        st.session_state.selected_currency_name = selected_currency_name # Seçilen para birimini de kaydet
        st.session_state.currency_symbol = currency_symbol # Sembolü de kaydet
        st.session_state.effective_exchange_rate = effective_exchange_rate # Efektif kuru da kaydet


    # Session state'den güncel değerleri al (hesaplamalar için)
    current_sales_input = st.session_state.sales_input_value
    current_smm_input = st.session_state.smm_input_value
    current_trade_receivables_input = st.session_state.trade_receivables_input_value
    current_inventories_input = st.session_state.inventories_input_value
    current_trade_payables_input = st.session_state.trade_payables_input_value
    current_current_assets_input = st.session_state.current_assets_input_value # Yeni
    current_current_liabilities_input = st.session_state.current_liabilities_input_value # Yeni
    current_selected_currency_name = st.session_state.selected_currency_name
    current_currency_symbol = st.session_state.currency_symbol
    current_exchange_rate_input = st.session_state.exchange_rate_input_value
    current_effective_exchange_rate = st.session_state.effective_exchange_rate


    with results_placeholder.container():
        st.header("⚙️ Calculation Results")

        # Giriş değerlerini TL'ye dönüştürme (eğer kur 1 değilse veya TL değilse)
        sales_for_calc = current_sales_input * current_effective_exchange_rate
        smm_for_calc = current_smm_input * current_effective_exchange_rate
        trade_receivables_for_calc = current_trade_receivables_input * current_effective_exchange_rate
        inventories_for_calc = current_inventories_input * current_effective_exchange_rate
        trade_payables_for_calc = current_trade_payables_input * current_effective_exchange_rate
        
        # Yeni eklenen manuel girişleri de TL'ye dönüştür
        current_assets_for_calc = current_current_assets_input * current_effective_exchange_rate
        current_liabilities_for_calc = current_current_liabilities_input * current_effective_exchange_rate

        # Kullanıcıya bilgilendirme: Eğer kur dönüşümü yapıldıysa
        if current_selected_currency_name != "TL" and current_exchange_rate_input != 1.0:
            st.info(f"Input values have been converted to TL (1 {current_selected_currency_name} = {current_exchange_rate_input:.4f} TL) for calculation purposes. Final financial amounts will be presented in {current_selected_currency_name} based on this conversion, while period calculations are in days.")
        elif current_selected_currency_name != "TL" and current_exchange_rate_input == 1.0:
             st.info(f"You selected {current_selected_currency_name} and entered an exchange rate of 1.0. Calculations are performed directly with your input values, and final financial amounts will be shown in {current_selected_currency_name}, while period calculations are in days.")
        else: # TL selected
             st.info(f"Calculations are performed with your input values in {current_selected_currency_name}. Financial amounts will be presented in {current_selected_currency_name}, and period calculations are in days.")


        if sales_for_calc == 0 or smm_for_calc == 0:
            st.warning("Annual Sales and Cost of Goods Sold (COGS) cannot be zero for calculation. Please enter valid values.")
            st.session_state.calculation_successful = False # Hesaplama başarısız oldu
        else:
            st.session_state.calculation_successful = True # Hesaplama başarılı oldu

            # Net Working Capital Cycle Calculations
            trade_receivable_collection_period = 0
            if trade_receivables_for_calc != 0:
                trade_receivable_collection_period = 360 / (sales_for_calc / trade_receivables_for_calc)
            else:
                st.info("Average Trade Receivables are zero, so Trade Receivable Collection Period is considered as 0.")

            inventory_holding_period = 0
            if inventories_for_calc != 0:
                inventory_holding_period = 360 / (smm_for_calc / inventories_for_calc)
            else:
                st.info("Average Inventories are zero, so Inventory Holding Period is considered as 0.")

            trade_payable_payment_period = 0
            if trade_payables_for_calc != 0:
                trade_payable_payment_period = 360 / (smm_for_calc / trade_payables_for_calc)
            else:
                st.info("Average Trade Payables are zero, so Trade Payable Payment Period is considered as 0.")

            net_working_capital_cycle = trade_receivable_collection_period + inventory_holding_period - trade_payable_payment_period

            st.subheader("1. Net Working Capital Cycle / Cash Conversion Cycle")
            col_metrics_1, col_metrics_2, col_metrics_3 = st.columns(3)
            with col_metrics_1:
                st.metric(label="Trade Receivable Collection Period", value=f"{trade_receivable_collection_period:.2f} days")
            with col_metrics_2:
                st.metric(label="Inventory Holding Period", value=f"{inventory_holding_period:.2f} days")
            with col_metrics_3:
                st.metric(label="Trade Payable Payment Period", value=f"{trade_payable_payment_period:.2f} days")
            st.metric(label="📊 NET WORKING CAPITAL CYCLE (Cash Conversion Cycle)", value=f"{net_working_capital_cycle:.2f} days")
            st.markdown("---")

            # --- MURAT'IN İSTEĞİ: MEVCUT NET İŞLETME SERMAYESİ VE İHTİYAÇ DUYULAN KISIM ---
            st.subheader("2. Net Working Capital (Current vs. Required)")

            # Mevcut Net İşletme Sermayesi Hesaplaması
            existing_net_working_capital_tl = current_assets_for_calc - current_liabilities_for_calc
            displayed_existing_nwc = existing_net_working_capital_tl / current_effective_exchange_rate

            col_nwc_1, col_nwc_2 = st.columns(2)
            with col_nwc_1:
                st.metric(label=f"Current Assets (Input)", value=f"{current_current_assets_input:,.2f} {current_currency_symbol}")
                st.metric(label=f"Current Liabilities (Input)", value=f"{current_current_liabilities_input:,.2f} {current_currency_symbol}")
                st.metric(label=f"💰 EXISTING NET WORKING CAPITAL", value=f"{displayed_existing_nwc:,.2f} {current_currency_symbol}", delta_color="off")
            with col_nwc_2:
                st.metric(label=f"Annual Sales", value=f"{current_sales_input:,.2f} {current_currency_symbol}")
                
                net_capital_duration_period = 0.0
                if existing_net_working_capital_tl > 0: # Mevcut net işletme sermayesi pozitifse hesapla
                    net_capital_duration_period = (existing_net_working_capital_tl / sales_for_calc) * 365
                else:
                    st.info("Existing Net Working Capital is zero or negative, so Net Capital Duration Period is not directly applicable in this context.")

                st.metric(label=f"Net Capital Duration Period", value=f"{net_capital_duration_period:.2f} days")
            st.markdown("---")


            # İhtiyaç Duyulan İşletme Sermayesi (Resimdeki "İhtiyaç Duyulan Net İşletme Sermayesi" satırı)
            required_nwc_based_on_cycle_tl = 0.0 # TL bazında
            if net_working_capital_cycle > 0:
                required_nwc_based_on_cycle_tl = (sales_for_calc / 365) * net_working_capital_cycle
            
            # Seçilen para birimine dönüştürülmüş hali
            displayed_required_nwc_based_on_cycle = required_nwc_based_on_cycle_tl / current_effective_exchange_rate
            
            # TOPLAM İHTİYAÇ OLAN NET İŞLETME SERMAYESİ (Senin daha önceki hesapladığın)
            st.subheader("3. TOTAL REQUIRED NET WORKING CAPITAL (from Cash Conversion Cycle)")
            
            required_nwc_amount_tl = 0.0 # TL bazında hesaplanan tutar
            if net_working_capital_cycle < 0:
                st.success("Congratulations! Your business has a positive cash conversion cycle. You do not require additional net working capital.")
                required_nwc_amount_tl = 0.0
            else:
                required_nwc_amount_tl = (sales_for_calc / 365) * net_working_capital_cycle
            
            # Hesaplanan TL tutarını seçilen para birimine geri dönüştür
            displayed_total_required_nwc = required_nwc_amount_tl / current_effective_exchange_rate

            st.metric(label=f"💰 TOTAL REQUIRED NET WORKING CAPITAL / CREDIT AMOUNT", value=f"{displayed_total_required_nwc:,.2f} {current_currency_symbol}")
            
            # İhtiyaç Duyulan Ek Sermaye (Murat'ın İsteği)
            additional_capital_needed = displayed_total_required_nwc - displayed_existing_nwc
            
            st.metric(
                label=f"ADDITIONAL CAPITAL REQUIRED /",
                value=f"{additional_capital_needed:,.2f} {current_currency_symbol}",
                delta=f"{additional_capital_needed:,.2f} {current_currency_symbol}" if additional_capital_needed != 0 else None,
                delta_color="inverse" if additional_capital_needed > 0 else "normal" if additional_capital_needed < 0 else "off"
            )

            st.markdown("---")

            # Hesaplanan tüm değerleri session state'e kaydet
            st.session_state.calculated_data = {
                'selected_currency_name': current_selected_currency_name,
                'currency_symbol': current_currency_symbol,
                'exchange_rate_input': current_exchange_rate_input,
                'effective_exchange_rate': current_effective_exchange_rate,
                'sales_input': current_sales_input,
                'smm_input': current_smm_input,
                'trade_receivables_input': current_trade_receivables_input,
                'inventories_input': current_inventories_input,
                'trade_payables_input': current_trade_payables_input,
                'current_assets_input': current_current_assets_input, # Yeni
                'current_liabilities_input': current_current_liabilities_input, # Yeni
                'sales_for_calc': sales_for_calc,
                'smm_for_calc': smm_for_calc,
                'trade_receivables_for_calc': trade_receivables_for_calc,
                'inventories_for_calc': inventories_for_calc,
                'trade_payables_for_calc': trade_payables_for_calc,
                'current_assets_for_calc': current_assets_for_calc, # Yeni
                'current_liabilities_for_calc': current_liabilities_for_calc, # Yeni
                'trade_receivable_collection_period': trade_receivable_collection_period,
                'inventory_holding_period': inventory_holding_period,
                'trade_payable_payment_period': trade_payable_payment_period,
                'net_working_capital_cycle': net_working_capital_cycle,
                'existing_net_working_capital_tl': existing_net_working_capital_tl, # Yeni
                'displayed_existing_nwc': displayed_existing_nwc, # Yeni
                'net_capital_duration_period': net_capital_duration_period, # Yeni
                'required_nwc_based_on_cycle_tl': required_nwc_based_on_cycle_tl, # Yeni
                'displayed_required_nwc_based_on_cycle': displayed_required_nwc_based_on_cycle, # Yeni
                'required_nwc_amount_tl': required_nwc_amount_tl, # Bu zaten vardı (Total Required)
                'displayed_total_required_nwc': displayed_total_required_nwc, # Bu da vardı (Total Required)
                'additional_capital_needed': additional_capital_needed # Yeni eklendi
            }

    # --- DOWNLOAD OPTIONS ---
    # Sadece hesaplama başarılıysa indirme seçeneklerini göster
    if st.session_state.calculation_successful:
        with download_placeholder.container():
            st.header("💾 Download Results")
            
            # Excel ve Word çıktıları için gerekli verileri DataFrame'e dönüştürelim
            # Burada tüm değerleri TL bazında gösteriyoruz, çünkü Excel'de ve Word'de daha tutarlı bir raporlama sağlar.
            # Kullanıcıya seçilen para birimini ve kuru en başta belirtiyoruz.
            data_for_export = {
                "Metric": [
                    "Selected Currency (Input)", "Exchange Rate (1 {} = ? TL)".format(st.session_state.calculated_data['selected_currency_name']),
                    "Annual Sales (Input Value)", "Cost of Goods Sold (COGS) (Input Value)", "Average Trade Receivables (Input Value)",
                    "Average Inventories (Input Value)", "Average Trade Payables (Input Value)",
                    "Current Assets (Input Value)", "Current Liabilities (Input Value)", 
                    "Annual Sales (Converted TL Value)", "Cost of Goods Sold (COGS) (Converted TL Value)", "Average Trade Receivables (Converted TL Value)",
                    "Average Inventories (Converted TL Value)", "Average Trade Payables (Converted TL Value)",
                    "Current Assets (Converted TL Value)", "Current Liabilities (Converted TL Value)",
                    "Trade Receivable Collection Period", "Inventory Holding Period",
                    "Trade Payable Payment Period", "NET WORKING CAPITAL CYCLE",
                    "EXISTING NET WORKING CAPITAL", 
                    "Net Capital Duration Period", 
                    "Required Net Working Capital (based on cycle)", 
                    "TOTAL REQUIRED NET WORKING CAPITAL (from Cash Conversion Cycle)",
                    "ADDITIONAL CAPITAL REQUIRED /" # Yeni eklendi
                ],
                "Value": [
                    st.session_state.calculated_data['selected_currency_name'], st.session_state.calculated_data['exchange_rate_input'],
                    st.session_state.calculated_data['sales_input'], st.session_state.calculated_data['smm_input'], st.session_state.calculated_data['trade_receivables_input'],
                    st.session_state.calculated_data['inventories_input'], st.session_state.calculated_data['trade_payables_input'],
                    st.session_state.calculated_data['current_assets_input'], st.session_state.calculated_data['current_liabilities_input'], 
                    st.session_state.calculated_data['sales_for_calc'], st.session_state.calculated_data['smm_for_calc'], st.session_state.calculated_data['trade_receivables_for_calc'],
                    st.session_state.calculated_data['inventories_for_calc'], st.session_state.calculated_data['trade_payables_for_calc'],
                    st.session_state.calculated_data['current_assets_for_calc'], st.session_state.calculated_data['current_liabilities_for_calc'], 
                    st.session_state.calculated_data['trade_receivable_collection_period'], st.session_state.calculated_data['inventory_holding_period'],
                    st.session_state.calculated_data['trade_payable_payment_period'], st.session_state.calculated_data['net_working_capital_cycle'],
                    st.session_state.calculated_data['displayed_existing_nwc'], 
                    st.session_state.calculated_data['net_capital_duration_period'], 
                    st.session_state.calculated_data['displayed_required_nwc_based_on_cycle'], 
                    st.session_state.calculated_data['displayed_total_required_nwc'],
                    st.session_state.calculated_data['additional_capital_needed'] # Yeni eklendi
                ],
                "Unit": [
                    "", "TL",
                    st.session_state.calculated_data['currency_symbol'], st.session_state.calculated_data['currency_symbol'], st.session_state.calculated_data['currency_symbol'],
                    st.session_state.calculated_data['currency_symbol'], st.session_state.calculated_data['currency_symbol'],
                    st.session_state.calculated_data['currency_symbol'], st.session_state.calculated_data['currency_symbol'], 
                    "TL", "TL", "TL",
                    "TL", "TL",
                    "TL", "TL", 
                    "days", "days", "days", "days", st.session_state.calculated_data['currency_symbol'], "days", st.session_state.calculated_data['currency_symbol'], st.session_state.calculated_data['currency_symbol'], st.session_state.calculated_data['currency_symbol'] # Yeni eklendi
                ]
            }
            df_export = pd.DataFrame(data_for_export)

            # --- Excel Çıktısı ---
            excel_file = BytesIO()
            with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
                df_export.to_excel(writer, sheet_name='NWC Analysis', index=False)
                workbook = writer.book
                worksheet = writer.sheets['NWC Analysis']
                
                for col_idx, col_name in enumerate(df_export.columns):
                    max_len = max(df_export[col_name].astype(str).map(len).max(), len(col_name)) + 2 
                    worksheet.column_dimensions[chr(65 + col_idx)].width = max_len
                    
                    if col_name == "Value":
                        for row_idx in range(2, len(df_export) + 2):
                            cell = worksheet.cell(row=row_idx, column=col_idx + 1)
                            # Apply currency format only for relevant rows
                            metric_name = df_export.iloc[row_idx - 2]["Metric"] # -2 because Excel rows start from 1, and header is row 1
                            
                            # Check if the metric name indicates it's a currency value in the selected currency
                            if any(x in metric_name for x in ["(Input Value)", "EXISTING NET WORKING CAPITAL", "Required Net Working Capital (based on cycle)", "TOTAL REQUIRED NET WORKING CAPITAL", "ADDITIONAL CAPITAL REQUIRED /"]): # "ADDITIONAL CAPITAL REQUIRED /" da eklendi
                                # Use the specific currency symbol for formatting
                                currency_format_str = f'#,##0.00 "{current_currency_symbol}"' if cell.value % 1 != 0 else f'#,##0 "{current_currency_symbol}"'
                                cell.number_format = currency_format_str
                            elif "Converted TL Value" in metric_name: # Ensure TL values are formatted with TL symbol
                                cell.number_format = '#,##0.00 "₺"' if cell.value % 1 != 0 else '#,##0 "₺"'
                            elif isinstance(cell.value, (int, float)): # Default numeric format for other numbers (like exchange rate)
                                cell.number_format = '#,##0.00' if cell.value % 1 != 0 else '#,##0'
                            
                            cell.alignment = Alignment(horizontal='right')
                
                for cell in worksheet["1:1"]:
                    cell.font = Font(bold=True)

            st.download_button(
                label="Download Excel Report",
                data=excel_file.getvalue(),
                file_name=f"Net_Working_Capital_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.document",
                help="Downloads the calculation results in an Excel file."
            )

            # --- Word Çıktısı (.docx) ---
            def create_word_report_from_session_state(data):
                document = Document()
                document.add_heading('Net Working Capital Analysis Report', level=1)
                document.add_paragraph(f"Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                document.add_paragraph("---")

                document.add_heading('Input Information', level=2)
                
                # Currency Information Table
                document.add_heading('Currency Information', level=3)
                table_curr = document.add_table(rows=2, cols=2)
                table_curr.style = 'Table Grid'
                table_curr.cell(0,0).text = "Selected Currency"
                table_curr.cell(0,1).text = "Exchange Rate (to TL)"
                table_curr.cell(1,0).text = data['selected_currency_name']
                table_curr.cell(1,1).text = f"1 {data['selected_currency_name']} = {data['exchange_rate_input']:,.4f} TL"

                for row in table_curr.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(10)
                            if row.cells[0].text in ["Selected Currency", "Exchange Rate (to TL)"]:
                                 run.bold = True

                # Input Values Table (in selected currency)
                document.add_heading(f"Input Values ({data['currency_symbol']})", level=3)
                input_data_rows = [
                    ["Metric", f"Value ({data['currency_symbol']})"],
                    ["Annual Sales", f"{data['sales_input']:,.2f} {data['currency_symbol']}"],
                    ["Cost of Goods Sold (COGS)", f"{data['smm_input']:,.2f} {data['currency_symbol']}"],
                    ["Average Trade Receivables", f"{data['trade_receivables_input']:,.2f} {data['currency_symbol']}"],
                    ["Average Inventories", f"{data['inventories_input']:,.2f} {data['currency_symbol']}"],
                    ["Average Trade Payables", f"{data['trade_payables_input']:,.2f} {data['currency_symbol']}"],
                    ["Current Assets", f"{data['current_assets_input']:,.2f} {data['currency_symbol']}"],
                    ["Current Liabilities", f"{data['current_liabilities_input']:,.2f} {data['currency_symbol']}"]
                ]
                
                table_input = document.add_table(rows=len(input_data_rows), cols=2)
                table_input.style = 'Table Grid'
                for r_idx, row_data in enumerate(input_data_rows):
                    for c_idx, cell_data in enumerate(row_data):
                        cell = table_input.cell(r_idx, c_idx)
                        cell.text = cell_data
                        if c_idx == 1 and r_idx > 0:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(10)
                            if r_idx == 0:
                                run.bold = True

                document.add_paragraph("---")
                document.add_heading('Calculation Results', level=2)

                # Converted Values (TL) Table - For internal calculation clarity
                if data['effective_exchange_rate'] != 1.0:
                    document.add_heading('Converted Values (TL for Calculation)', level=3)
                    converted_data_rows = [
                        ["Metric", "Value (TL)"],
                        ["Annual Sales", f"{data['sales_for_calc']:,.2f} TL"],
                        ["Cost of Goods Sold (COGS)", f"{data['smm_for_calc']:,.2f} TL"],
                        ["Average Trade Receivables", f"{data['trade_receivables_for_calc']:,.2f} TL"],
                        ["Average Inventories", f"{data['inventories_for_calc']:,.2f} TL"],
                        ["Average Trade Payables", f"{data['trade_payables_for_calc']:,.2f} TL"],
                        ["Current Assets", f"{data['current_assets_for_calc']:,.2f} TL"],
                        ["Current Liabilities", f"{data['current_liabilities_for_calc']:,.2f} TL"]
                    ]
                    table_converted = document.add_table(rows=len(converted_data_rows), cols=2)
                    table_converted.style = 'Table Grid'
                    for r_idx, row_data in enumerate(converted_data_rows):
                        for c_idx, cell_data in enumerate(row_data):
                            cell = table_converted.cell(r_idx, c_idx)
                            cell.text = cell_data
                            if c_idx == 1 and r_idx > 0:
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            for run in cell.paragraphs[0].runs:
                                run.font.size = Pt(10)
                                if r_idx == 0:
                                    run.bold = True
                
                document.add_heading('Net Working Capital Cycle / Cash Conversion Cycle', level=3)
                nwc_cycle_data_rows = [
                    ["Metric", "Value (days)"],
                    ["Trade Receivable Collection Period", f"{data['trade_receivable_collection_period']:,.2f}"],
                    ["Inventory Holding Period", f"{data['inventory_holding_period']:,.2f}"],
                    ["Trade Payable Payment Period", f"{data['trade_payable_payment_period']:,.2f}"],
                    ["NET WORKING CAPITAL CYCLE", f"{data['net_working_capital_cycle']:,.2f}"]
                ]
                table_nwc_cycle = document.add_table(rows=len(nwc_cycle_data_rows), cols=2)
                table_nwc_cycle.style = 'Table Grid'
                for r_idx, row_data in enumerate(nwc_cycle_data_rows):
                    for c_idx, cell_data in enumerate(row_data):
                        cell = table_nwc_cycle.cell(r_idx, c_idx)
                        cell.text = cell_data
                        if c_idx == 1 and r_idx > 0:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(10)
                            if r_idx == 0 or "NET WORKING CAPITAL CYCLE" in cell_data:
                                run.bold = True

                document.add_heading('Net Working Capital (Current vs. Required)', level=3)
                existing_nwc_rows = [
                    ["Metric", f"Value ({data['currency_symbol']})"],
                    ["Current Assets (Input)", f"{data['current_assets_input']:,.2f} {data['currency_symbol']}"],
                    ["Current Liabilities (Input)", f"{data['current_liabilities_input']:,.2f} {data['currency_symbol']}"],
                    ["EXISTING NET WORKING CAPITAL", f"{data['displayed_existing_nwc']:,.2f} {data['currency_symbol']}"],
                    ["Annual Sales (from Input)", f"{data['sales_input']:,.2f} {data['currency_symbol']}"],
                    ["Net Capital Duration Period", f"{data['net_capital_duration_period']:,.2f} days"],
                    # "Required Net Working Capital (based on cycle)" satırı buradan kaldırıldı.
                    ["TOTAL REQUIRED NET WORKING CAPITAL (from Cycle)", f"{data['displayed_total_required_nwc']:,.2f} {data['currency_symbol']}"], # Eski TOTAL REQUIRED'ı buraya taşıdık
                    ["🎯 ADDITIONAL CAPITAL REQUIRED /", f"{data['additional_capital_needed']:,.2f} {data['currency_symbol']}"] # Yeni satır
                ]
                table_existing_nwc = document.add_table(rows=len(existing_nwc_rows), cols=2)
                table_existing_nwc.style = 'Table Grid'
                for r_idx, row_data in enumerate(existing_nwc_rows):
                    for c_idx, cell_data in enumerate(row_data):
                        cell = table_existing_nwc.cell(r_idx, c_idx)
                        cell.text = cell_data
                        if c_idx == 1 and r_idx > 0:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(10)
                            if r_idx == 0 or "EXISTING NET WORKING CAPITAL" in cell_data or "Net Capital Duration Period" in cell_data or "TOTAL REQUIRED NET WORKING CAPITAL" in cell_data or "ADDITIONAL CAPITAL REQUIRED" in cell_data:
                                run.bold = True
                
                # TOTAL REQUIRED NET WORKING CAPITAL (from Cash Conversion Cycle) kısmı artık yukarıdaki tabloda bir satır olarak yer alıyor.
                # document.add_heading('TOTAL REQUIRED NET WORKING CAPITAL (from Cash Conversion Cycle)', level=3)
                # total_required_nwc_rows = [
                #     ["Metric", f"Value ({data['currency_symbol']})"],
                #     ["TOTAL REQUIRED NET WORKING CAPITAL / CREDIT AMOUNT", f"{data['displayed_total_required_nwc']:,.2f} {data['currency_symbol']}"]
                # ]
                # table_total_required_nwc = document.add_table(rows=len(total_required_nwc_rows), cols=2)
                # table_total_required_nwc.style = 'Table Grid'
                # for r_idx, row_data in enumerate(total_required_nwc_rows):
                #     for c_idx, cell_data in enumerate(row_data):
                #         cell = table_total_required_nwc.cell(r_idx, c_idx)
                #         cell.text = cell_data
                #         if c_idx == 1 and r_idx > 0:
                #             cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                #         else:
                #             cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                #         cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                #         for run in cell.paragraphs[0].runs:
                #             run.font.size = Pt(10)
                #             if r_idx == 0 or "TOTAL REQUIRED NET WORKING CAPITAL" in cell_data:
                #                 run.bold = True

                document.add_paragraph("---")
                document.add_heading('Information', level=2)
                document.add_paragraph("Net Working Capital = Current Assets - Current Liabilities")
                document.add_paragraph("This report differentiates between 'Existing Net Working Capital' (calculated from Current Assets and Current Liabilities) and 'Required Net Working Capital' (calculated based on the cash conversion cycle). The 'TOTAL REQUIRED NET WORKING CAPITAL' indicates the amount of funding needed to support the operational cycle of the business. 'ADDITIONAL CAPITAL REQUIRED /' shows the net difference between total required capital and your existing working capital. A positive value indicates additional funding needed, while a negative value indicates a surplus.")

                doc_buffer = BytesIO()
                document.save(doc_buffer)
                doc_buffer.seek(0)
                return doc_buffer

            word_file = create_word_report_from_session_state(st.session_state.calculated_data)

            st.download_button(
                label="Download Word Report (.docx)",
                data=word_file.getvalue(),
                file_name=f"Net_Working_Capital_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Downloads the calculation results in a Word document (.docx)."
            )