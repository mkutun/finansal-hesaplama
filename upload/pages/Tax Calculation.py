import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches

def calculate_sahis_vergisi(gelir, texts):
    """
    Şahıs şirketi gelir vergisini hesaplar.
    2025 yılı vergi dilimlerine göre hesaplama yapar ve detayları döndürür.
    """
    vergi = 0
    hesaplama_detaylari = []

    # İlk dilim
    if gelir <= 158000:
        vergi_bu_dilim = gelir * 0.15
        vergi += vergi_bu_dilim
        hesaplama_detaylari.append(f"• ₺{gelir:,.2f} {texts['income_of']} %15'i ({texts['tax_bracket_label']} {texts['up_to']} ₺158.000): ₺{vergi_bu_dilim:,.2f}")
        dilim = "15%"
    else:
        vergi_bu_dilim = 158000 * 0.15
        vergi += vergi_bu_dilim
        hesaplama_detaylari.append(f"• ₺158.000 {texts['income_of']} %15'i ({texts['tax_bracket_label']} {texts['up_to']} ₺158.000): ₺{vergi_bu_dilim:,.2f}")
        kalan_gelir = gelir - 158000

        # İkinci dilim
        if kalan_gelir > 0:
            if kalan_gelir <= (330000 - 158000):
                vergi_bu_dilim = kalan_gelir * 0.20
                vergi += vergi_bu_dilim
                hesaplama_detaylari.append(f"• ₺{kalan_gelir:,.2f} {texts['income_of']} %20'si ({texts['tax_bracket_label']} ₺158.000-₺330.000 {texts['between']}): ₺{vergi_bu_dilim:,.2f}")
                dilim = "20%"
            else:
                vergi_bu_dilim = (330000 - 158000) * 0.20
                vergi += vergi_bu_dilim
                hesaplama_detaylari.append(f"• ₺{(330000 - 158000):,.2f} {texts['income_of']} %20'si ({texts['tax_bracket_label']} ₺158.000-₺330.000 {texts['between']}): ₺{vergi_bu_dilim:,.2f}")
                kalan_gelir -= (330000 - 158000)
                dilim = "20%"

                # Üçüncü dilim
                if kalan_gelir > 0:
                    if kalan_gelir <= (800000 - 330000):
                        vergi_bu_dilim = kalan_gelir * 0.27
                        vergi += vergi_bu_dilim
                        hesaplama_detaylari.append(f"• ₺{kalan_gelir:,.2f} {texts['income_of']} %27'si ({texts['tax_bracket_label']} ₺330.000-₺800.000 {texts['between']}): ₺{vergi_bu_dilim:,.2f}")
                        dilim = "27%"
                    else:
                        vergi_bu_dilim = (800000 - 330000) * 0.27
                        vergi += vergi_bu_dilim
                        hesaplama_detaylari.append(f"• ₺{(800000 - 330000):,.2f} {texts['income_of']} %27'si ({texts['tax_bracket_label']} ₺330.000-₺800.000 {texts['between']}): ₺{vergi_bu_dilim:,.2f}")
                        kalan_gelir -= (800000 - 330000)
                        dilim = "27%"

                        # Dördüncü dilim
                        if kalan_gelir > 0:
                            if kalan_gelir <= (4300000 - 800000):
                                vergi_bu_dilim = kalan_gelir * 0.35
                                vergi += vergi_bu_dilim
                                hesaplama_detaylari.append(f"• ₺{kalan_gelir:,.2f} {texts['income_of']} %35'i ({texts['tax_bracket_label']} ₺800.000-₺4.300.000 {texts['between']}): ₺{vergi_bu_dilim:,.2f}")
                                dilim = "35%"
                            else:
                                vergi_bu_dilim = (4300000 - 800000) * 0.35
                                vergi += vergi_bu_dilim
                                hesaplama_detaylari.append(f"• ₺{(4300000 - 800000):,.2f} {texts['income_of']} %35'i ({texts['tax_bracket_label']} ₺800.000-₺4.300.000 {texts['between']}): ₺{vergi_bu_dilim:,.2f}")
                                kalan_gelir -= (4300000 - 800000)
                                dilim = "35%"

                                # Beşinci dilim (ve sonrası)
                                if kalan_gelir > 0:
                                    vergi_bu_dilim = kalan_gelir * 0.40
                                    vergi += vergi_bu_dilim
                                    hesaplama_detaylari.append(f"• ₺{kalan_gelir:,.2f} {texts['income_of']} %40'ı ({texts['tax_bracket_label']} ₺4.300.000 {texts['above']}): ₺{vergi_bu_dilim:,.2f}")
                                    dilim = "40%"

    return vergi, dilim, hesaplama_detaylari

def get_output_text(lang):
    """
    Seçilen dile göre çıktı metinlerini döndürür.
    """
    if lang == "Türkçe":
        texts = {
            "title": "Vergi Hesaplama Uygulaması",
            "company_type": "Firma Tipi Seçimi:",
            "sahis": "Şahıs Şirketi",
            "ltd": "Limited Şirket",
            "income_input": "Yıllık Gelirinizi Girin (₺):",
            "calculate_button": "Vergiyi Hesapla",
            "tax_info_sahis_title": "Şahıs Şirketi Vergi Bilgileri",
            "tax_info_ltd_title": "Limited Şirket Vergi Bilgileri",
            "income_label": "Yıllık Gelir",
            "calculated_tax_label": "Hesaplanan Vergi",
            "tax_bracket_label": "Vergi Dilimi",
            "corporate_tax_rate_label": "Kurumlar Vergisi Oranı",
            "corporate_tax_label": "Kurumlar Vergisi Tutarı",
            "download_excel": "Excel Olarak İndir",
            "download_word": "Word Olarak İndir",
            "error_invalid_input": "Lütfen geçerli bir yıllık gelir tutarı girin.",
            "sahis_desc": "Şahıs şirketleri, artan oranlı gelir vergisi dilimlerine tabidir. Detaylı hesaplama aşağıdaki gibidir:",
            "ltd_desc": "Limited şirketler, %25 oranında kurumlar vergisine tabidir. Hesaplama detayları aşağıdaki gibidir:",
            "language_selection": "Dil Seçimi:",
            "up_to": "kadar",
            "between": "arası",
            "above": "üzeri",
            "income_of": "gelirin",
            "total_tax": "Toplam Hesaplanan Vergi",
            "current": "güncel",
            "tax_calculation_details_title": "Vergi Hesaplama Detayları", # Bu satırı ekledik
            "summary_table_title": "Özet Tablo" # Bu satırı ekledik
        }
    else: # English
        texts = {
            "title": "Tax Calculation Application",
            "company_type": "Select Company Type:",
            "sahis": "Sole Proprietorship",
            "ltd": "Limited Company",
            "income_input": "Enter Your Annual Income (₺):",
            "calculate_button": "Calculate Tax",
            "tax_info_sahis_title": "Sole Proprietorship Tax Information",
            "tax_info_ltd_title": "Limited Company Tax Information",
            "income_label": "Annual Income",
            "calculated_tax_label": "Calculated Tax",
            "tax_bracket_label": "Tax Bracket",
            "corporate_tax_rate_label": "Corporate Tax Rate",
            "corporate_tax_label": "Corporate Tax Amount",
            "download_excel": "Download as Excel",
            "download_word": "Download as Word",
            "error_invalid_input": "Please enter a valid annual income.",
            "sahis_desc": "Sole proprietorships are subject to progressive income tax rates. Detailed calculation is as follows:",
            "ltd_desc": "Limited companies are subject to a 25% corporate tax rate. Calculation details are as follows:",
            "language_selection": "Language Selection:",
            "up_to": "up to",
            "between": "between",
            "above": "above",
            "income_of": "income of",
            "total_tax": "Total Calculated Tax",
            "current": "current",
            "tax_calculation_details_title": "Tax Calculation Details", # Bu satırı ekledik
            "summary_table_title": "Summary Table" # Bu satırı ekledik
        }
    return texts

def create_word_document(df, title, company_type, lang_texts, details=None):
    """
    Pandas DataFrame'i Word belgesine dönüştürür.
    """
    document = Document()
    document.add_heading(title, level=1)

    # Genel Bilgiler
    document.add_paragraph(f"{lang_texts['company_type'].replace(':', '')}: {company_type}")
    document.add_paragraph(f"{lang_texts['income_label']}: {df[lang_texts['income_label']].iloc[0]}")

    # Detaylar (Şahıs Şirketi ve Limited Şirket için)
    document.add_heading(lang_texts["tax_calculation_details_title"], level=2)
    if company_type == lang_texts["sahis"] and details:
        for detail in details:
            document.add_paragraph(detail)
        document.add_paragraph(f"**{lang_texts['total_tax']}: {df[lang_texts['calculated_tax_label']].iloc[0]}**")
        document.add_paragraph(f"**{lang_texts['tax_bracket_label']} ({lang_texts['current']}): {df[lang_texts['tax_bracket_label']].iloc[0]}**")
    elif company_type == lang_texts["ltd"]:
        document.add_paragraph(f"• {lang_texts['income_label']}: {df[lang_texts['income_label']].iloc[0]}")
        document.add_paragraph(f"• {lang_texts['corporate_tax_rate_label']}: {df[lang_texts['corporate_tax_rate_label']].iloc[0]}")
        document.add_paragraph(f"• {lang_texts['corporate_tax_label']}: {df[lang_texts['corporate_tax_label']].iloc[0]}")


    # Tabloyu ekle
    document.add_heading(lang_texts["summary_table_title"], level=2)
    table = document.add_table(df.shape[0]+1, df.shape[1])
    # Word dokümanında sütun genişliğini dinamik yapmak yerine sabit bir değer belirleyelim
    # table.columns[0].width = Inches(2) # Bu satır bazen hata verebilir, kaldırdım

    # Başlıklar
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = col
    
    # Satırlar
    for i, row in df.iterrows():
        for j, cell in enumerate(row):
            table.cell(i+1, j).text = str(cell)

    document.add_page_break()

    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio

# Streamlit Uygulaması Başlangıcı
st.set_page_config(layout="wide") # Sayfa düzenini genişletir

# Dil seçimi
lang_options = ["Türkçe", "English"]
selected_lang = st.sidebar.selectbox("Dil Seçimi / Language Selection", lang_options)
texts = get_output_text(selected_lang) # Metinleri yükle

st.title(texts["title"])

st.markdown("---")

# Firma Tipi Seçimi
company_type = st.radio(
    texts["company_type"],
    (texts["sahis"], texts["ltd"]),
    horizontal=True
)

# Yıllık Gelir Girişi
st.markdown(f"### {texts['income_input']}")
gelir = st.number_input(texts["income_input"], min_value=0.0, format="%.2f", key="gelir_input", label_visibility="collapsed")

st.markdown("---")

if st.button(texts["calculate_button"]):
    if gelir is None or gelir == 0:
        st.error(texts["error_invalid_input"])
    else:
        st.success(f"{texts['calculated_tax_label']} Hesaplanıyor...")

        if company_type == texts["sahis"]:
            st.header(texts["tax_info_sahis_title"])
            st.info(texts["sahis_desc"])
            vergi, dilim, hesaplama_detaylari = calculate_sahis_vergisi(gelir, texts)

            # Hesaplama detaylarını liste olarak göster
            for detail in hesaplama_detaylari:
                st.markdown(detail)

            st.markdown(f"**{texts['total_tax']}: ₺{vergi:,.2f}**")
            st.markdown(f"**{texts['tax_bracket_label']} ({texts['current']}): {dilim}**")


            data = {
                texts["income_label"]: [f"₺{gelir:,.2f}"],
                texts["calculated_tax_label"]: [f"₺{vergi:,.2f}"],
                texts["tax_bracket_label"]: [dilim]
            }
            df = pd.DataFrame(data)
            st.subheader(texts["summary_table_title"]) # Özet tablo başlığı
            st.table(df)

            # Excel çıktısı
            output_excel = BytesIO()
            with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
                df.to_excel(writer, index=False, sheet_name='Vergi Bilgileri')
            output_excel.seek(0)
            st.download_button(
                label=texts["download_excel"],
                data=output_excel,
                file_name=f"{texts['sahis']}_Vergi_Hesaplama.xlsx" if selected_lang == "Türkçe" else "SoleProprietorship_Tax_Calculation.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

            # Word çıktısı
            word_output = create_word_document(
                df,
                texts["tax_info_sahis_title"],
                company_type,
                texts,
                details=hesaplama_detaylari
            )
            st.download_button(
                label=texts["download_word"],
                data=word_output,
                file_name=f"{texts['sahis']}_Vergi_Hesaplama.docx" if selected_lang == "Türkçe" else "SoleProprietorship_Tax_Calculation.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        elif company_type == texts["ltd"]:
            st.header(texts["tax_info_ltd_title"])
            st.info(texts["ltd_desc"])
            kurumlar_vergisi_orani = 0.25
            kurumlar_vergisi_tutari = gelir * kurumlar_vergisi_orani

            # Limited Şirket detaylarını ekrana yazdırma
            st.markdown(f"• {texts['income_label']}: ₺{gelir:,.2f}")
            st.markdown(f"• {texts['corporate_tax_rate_label']}: %{kurumlar_vergisi_orani*100:,.0f}")
            st.markdown(f"**{texts['corporate_tax_label']}: ₺{kurumlar_vergisi_tutari:,.2f}**")


            data = {
                texts["income_label"]: [f"₺{gelir:,.2f}"],
                texts["corporate_tax_rate_label"]: ["%25"],
                texts["corporate_tax_label"]: [f"₺{kurumlar_vergisi_tutari:,.2f}"]
            }
            df = pd.DataFrame(data)
            st.subheader(texts["summary_table_title"]) # Özet tablo başlığı
            st.table(df)

            # Excel çıktısı
            output_excel = BytesIO()
            with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
                df.to_excel(writer, index=False, sheet_name='Vergi Bilgileri')
            output_excel.seek(0)
            st.download_button(
                label=texts["download_excel"],
                data=output_excel,
                file_name=f"{texts['ltd']}_Vergi_Hesaplama.xlsx" if selected_lang == "Türkçe" else "LimitedCompany_Tax_Calculation.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

            # Word çıktısı
            word_output = create_word_document(
                df,
                texts["tax_info_ltd_title"],
                company_type,
                texts
            )
            st.download_button(
                label=texts["download_word"],
                data=word_output,
                file_name=f"{texts['ltd']}_Vergi_Hesaplama.docx" if selected_lang == "Türkçe" else "LimitedCompany_Tax_Calculation.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

st.markdown("---")
st.markdown("Uygulamanın 2025 yılı vergi dilimlerine göre hazırlandığını unutmayın. Yasal mali müşavirlik tavsiyesi değildir.")