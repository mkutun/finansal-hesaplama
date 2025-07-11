import streamlit as st

# Kullanıcının giriş yapıp yapmadığını kontrol et
if not st.session_state.get('logged_in', False):
    st.warning("Bu sayfayı görüntülemek için giriş yapmanız gerekmektedir.")
    st.switch_page("Home_Page.py") # Giriş sayfasına yönlendir
    st.stop() # Sayfanın geri kalan kodunu çalıştırmayı durdur
import streamlit as st
import pandas as pd
import io
import math
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
from openpyxl.styles.numbers import BUILTIN_FORMATS

# --- Constants and Settings ---
CURRENCY_SYMBOLS = {
    "₺": "TL",
    "$": "USD",
    "€": "EUR",
    "£": "GBP"
}
CURRENCY_NAMES = {v: k for k, v in CURRENCY_SYMBOLS.items()}

# --- Helper function for number formatting (for display in Streamlit) ---
def format_number(number, is_year=False, include_currency=True, currency_symbol="₺", is_percentage=False):
    """
    Formats a number.
    If is_year is True, formats as an integer without currency or decimals.
    If include_currency is False, it formats without the currency symbol.
    If is_percentage is True, formats as a percentage with one decimal place.
    Otherwise, formats with thousands commas and two decimal places (English locale style).
    """
    if pd.isna(number) or number == '':
        return ''

    if is_year:
        return f"{int(number)}"
    
    if is_percentage:
        return f"{number:.1f}%" # English style percentage
    
    # Format with thousands commas and two decimal places (English locale style)
    formatted_output = f"{number:,.2f}" 
    
    if include_currency:
        # Currency symbol after number for TL, otherwise depends on common usage.
        # Sticking to current code's placement for consistency unless specified.
        return f"{formatted_output} {currency_symbol}"
    else:
        return formatted_output


# --- Main calculation function ---
def calculate_loan_repayment_schedule(principal_amount, annual_interest_rate_percent, grace_period_years, total_loan_term_years):
    """
    Calculates the loan repayment schedule, including automatically computed annual payments.
    Returns the schedule data as a list of dictionaries and the total amount paid.
    """
    if not all(isinstance(val, (int, float)) for val in [principal_amount, annual_interest_rate_percent, grace_period_years, total_loan_term_years]):
        return [], 0.0

    annual_interest_rate_decimal = annual_interest_rate_percent / 100.0
    schedule_data = []
    current_balance = principal_amount
    start_year = pd.to_datetime('today').year

    # Grace Period Calculation
    for year_offset in range(grace_period_years):
        current_year = start_year + year_offset
        interest_accrued = current_balance * annual_interest_rate_decimal
        current_balance += interest_accrued
        schedule_data.append({
            'YEAR': current_year,
            'PRINCIPAL PAYMENT': 0.0,
            'INTEREST': interest_accrued,
            'P+I': interest_accrued,
            'PAYMENT': 0.0,
            'REMAINING BALANCE': current_balance
        })

    # Repayment Period Calculation
    repayment_start_year = start_year + grace_period_years
    repayment_period_years = total_loan_term_years - grace_period_years

    if repayment_period_years <= 0:
        st.error("Total loan term must be greater than the grace period. Please check the values.")
        return [], 0.0

    pmt_principal_for_repayment = current_balance
    
    if annual_interest_rate_decimal == 0:
        annual_payment = pmt_principal_for_repayment / repayment_period_years
    else:
        annual_payment = (pmt_principal_for_repayment * annual_interest_rate_decimal) / \
                         (1 - (1 + annual_interest_rate_decimal)**(-repayment_period_years))
    
    total_amount_paid = 0.0

    for i in range(repayment_period_years):
        current_year = repayment_start_year + i
        interest_for_this_year = current_balance * annual_interest_rate_decimal
        
        if i == repayment_period_years - 1:
            payment_this_year = current_balance + interest_for_this_year
            principal_payment_this_year = current_balance
            current_balance = 0.0
        else:
            payment_this_year = annual_payment
            principal_payment_this_year = payment_this_year - interest_for_this_year
            current_balance -= principal_payment_this_year
        
        if abs(current_balance) < 0.01:
            current_balance = 0.0
            
        principal_plus_interest = principal_payment_this_year + interest_for_this_year
        total_amount_paid += payment_this_year

        schedule_data.append({
            'YEAR': current_year,
            'PRINCIPAL PAYMENT': principal_payment_this_year,
            'INTEREST': interest_for_this_year,
            'P+I': principal_plus_interest,
            'PAYMENT': payment_this_year,
            'REMAINING BALANCE': current_balance
        })
    
    return schedule_data, total_amount_paid

# --- Report Generation Functions ---

def create_excel_xlsx_report(schedule_data, total_payment, principal_amount, interest_rate, grace_period, total_loan_term, currency_symbol):
    wb = Workbook()
    ws = wb.active
    ws.title = "Loan Repayment Schedule" # Excel sheet title English

    # Define styles
    header_font = Font(bold=True, size=10)
    data_font = Font(size=9)
    right_aligned = Alignment(horizontal='right')
    center_aligned = Alignment(horizontal='center')
    thin_border = Border(left=Side(style='thin'), 
                         right=Side(style='thin'), 
                         top=Side(style='thin'), 
                         bottom=Side(style='thin'))
    
    # Custom number format for currency (e.g., #,##0.00 "₺")
    currency_excel_format = f'#,##0.00 "{currency_symbol}"'
    
    # Add general information
    ws.append(["--- LOAN REPAYMENT SCHEDULE ---"]) # Title English
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=6)
    ws['A1'].alignment = center_aligned
    ws['A1'].font = header_font

    repayment_period_years = total_loan_term - grace_period
    ws.append([f"{grace_period} YEARS GRACE; {repayment_period_years} YEARS PAYMENT; TOTAL {total_loan_term} YEARS LOAN TERM"]) # Description English
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=6)
    ws['A2'].alignment = center_aligned
    ws['A2'].font = data_font

    ws.append([]) # Blank row
    ws.append(["Loan Principal:", principal_amount]) # Label English
    ws.cell(row=4, column=2).number_format = currency_excel_format
    
    # Interest rate row and alignment
    ws.append(["Annual Interest Rate:", f"%{interest_rate:.1f}"]) # Format updated and label English
    ws.cell(row=ws.max_row, column=2).alignment = right_aligned # Right align
    ws.append([]) # Blank row

    # Add table headers
    headers = ['YEAR', 'PRINCIPAL PAYMENT', 'INTEREST', 'P+I', 'PAYMENT', 'REMAINING BALANCE'] # Headers English
    ws.append(headers)

    for col_idx, header_text in enumerate(headers, 1):
        cell = ws.cell(row=ws.max_row, column=col_idx)
        cell.font = header_font
        cell.alignment = center_aligned
        cell.border = thin_border

    # Add data rows
    for row_data in schedule_data:
        row_values = [
            row_data['YEAR'],
            row_data['PRINCIPAL PAYMENT'],
            row_data['INTEREST'],
            row_data['P+I'],
            row_data['PAYMENT'],
            row_data['REMAINING BALANCE']
        ]
        ws.append(row_values)
        
        # Apply formatting to data rows
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=ws.max_row, column=col_idx)
            cell.font = data_font
            cell.border = thin_border
            if col_idx == 1: # YEAR column
                cell.alignment = center_aligned
                cell.number_format = BUILTIN_FORMATS[1] # Integer format
            else: # Numeric columns
                cell.alignment = right_aligned
                cell.number_format = currency_excel_format # Apply custom currency format

    ws.append([]) # Blank row
    ws.append(["Total Amount Paid:", "", "", "", "", total_payment]) # Label English
    ws.merge_cells(start_row=ws.max_row, start_column=1, end_row=ws.max_row, end_column=5)
    ws.cell(row=ws.max_row, column=1).alignment = right_aligned
    ws.cell(row=ws.max_row, column=6).number_format = currency_excel_format
    ws.cell(row=ws.max_row, column=6).alignment = right_aligned
    ws.cell(row=ws.max_row, column=1).font = header_font
    ws.cell(row=ws.max_row, column=6).font = header_font


    # Adjust column widths automatically
    for col_idx in range(1, ws.max_column + 1):
        max_length = 0
        for row_idx in range(1, ws.max_row + 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            if cell.value is not None:
                cell_value_str = str(cell.value)
                if len(cell_value_str) > max_length:
                    max_length = len(cell_value_str)
        
        column_letter = get_column_letter(col_idx)
        adjusted_width = (max_length + 2) * 1.2 # Add some padding
        if adjusted_width > 0:
            ws.column_dimensions[column_letter].width = adjusted_width

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.getvalue()

def create_word_docx_report(schedule_data, total_payment, principal_amount, interest_rate, grace_period, total_loan_term, currency_symbol):
    document = Document()

    document.add_heading('Loan Repayment Schedule', level=1) # Heading English
    
    repayment_period_years = total_loan_term - grace_period
    document.add_paragraph(f"{grace_period} Years Grace, {repayment_period_years} Years Payment, Total {total_loan_term} Years Loan Term") # Description English
    document.add_paragraph(f"Loan Principal: {format_number(principal_amount, currency_symbol=currency_symbol)}") # Label English
    document.add_paragraph(f"Annual Interest Rate: {format_number(interest_rate, is_percentage=True)}") # Label English and percentage format used
    document.add_paragraph("\n")

    headers = ['YEAR', 'PRINCIPAL PAYMENT', 'INTEREST', 'P+I', 'PAYMENT', 'REMAINING BALANCE'] # Headers English
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    for row in schedule_data:
        row_cells = table.add_row().cells
        row_cells[0].text = format_number(row['YEAR'], is_year=True)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        row_cells[1].text = format_number(row['PRINCIPAL PAYMENT'], currency_symbol=currency_symbol)
        row_cells[2].text = format_number(row['INTEREST'], currency_symbol=currency_symbol)
        row_cells[3].text = format_number(row['P+I'], currency_symbol)
        row_cells[4].text = format_number(row['PAYMENT'], currency_symbol=currency_symbol)
        row_cells[5].text = format_number(row['REMAINING BALANCE'], currency_symbol=currency_symbol)

        for j in range(1, len(headers)):
            row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    document.add_paragraph("\n")
    document.add_paragraph(f"Total Amount Paid: {format_number(total_payment, currency_symbol=currency_symbol)}") # Label English

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()


# --- Streamlit Application Layout ---
st.set_page_config(layout="wide", page_title="Loan Calculation Wizard") # Title English

st.title("Loan Calculation Wizard") # Title English
st.markdown("Enter loan parameters and view the repayment schedule and total amount paid.") # Description English

# --- Session State Initialization ---
if 'schedule_data' not in st.session_state:
    st.session_state.schedule_data = []
if 'total_payment' not in st.session_state:
    st.session_state.total_payment = 0.0
if 'show_results' not in st.session_state:
    st.session_state.show_results = False
if 'principal_amount' not in st.session_state:
    st.session_state.principal_amount = 450000000.0
if 'interest_rate' not in st.session_state:
    st.session_state.interest_rate = 3.5
if 'grace_period' not in st.session_state:
    st.session_state.grace_period = 3
if 'total_loan_term' not in st.session_state:
    st.session_state.total_loan_term = 18
if 'selected_currency_symbol' not in st.session_state:
    st.session_state.selected_currency_symbol = "₺"


# --- Input Section ---
st.header("Loan Parameters") # Header English

col1, col2 = st.columns(2)

with col1:
    st.session_state.principal_amount = st.number_input(
        "Principal Amount:", # Label English
        min_value=0.0,
        value=st.session_state.principal_amount,
        step=1000.0,
        format="%.0f", # Integer format for principal
        key="principal_input"
    )
    st.session_state.interest_rate = st.number_input(
        " Annual Interest Rate (%):", # Label English
        min_value=0.0,
        max_value=100.0,
        value=st.session_state.interest_rate,
        step=0.1,
        format="%.1f", # One decimal place for interest rate
        key="interest_input"
    )

with col2:
    st.session_state.grace_period = st.number_input(
        "Grace Period (Years):", # Label English
        min_value=0,
        value=st.session_state.grace_period,
        step=1,
        key="grace_input"
    )
    st.session_state.total_loan_term = st.number_input(
        "Total Loan Term (Years):", # Label English
        min_value=1,
        value=st.session_state.total_loan_term,
        step=1,
        key="term_input"
    )
    st.session_state.selected_currency_symbol = st.selectbox(
        "Currency:", # Label English
        options=list(CURRENCY_SYMBOLS.keys()),
        index=list(CURRENCY_SYMBOLS.keys()).index(st.session_state.selected_currency_symbol) if st.session_state.selected_currency_symbol in CURRENCY_SYMBOLS else 0,
        key="currency_select"
    )

# --- Calculate Button ---
if st.button("Calculate and Show Results", key="calculate_btn"): # Button English
    if st.session_state.principal_amount <= 0 or st.session_state.interest_rate < 0 or st.session_state.grace_period < 0 or st.session_state.total_loan_term <= 0:
        st.error("Please enter valid positive numbers for all fields.") # Error message English
        st.session_state.show_results = False
    elif st.session_state.grace_period >= st.session_state.total_loan_term:
        st.error("Grace period must be less than the total loan term.") # Error message English
        st.session_state.show_results = False
    else:
        st.session_state.schedule_data, st.session_state.total_payment = calculate_loan_repayment_schedule(
            st.session_state.principal_amount,
            st.session_state.interest_rate,
            st.session_state.grace_period,
            st.session_state.total_loan_term
        )
        if st.session_state.schedule_data:
            st.session_state.show_results = True
        else:
            st.session_state.show_results = False


# --- Display Results (controlled by session_state) ---
if st.session_state.show_results:
    st.header("Loan Repayment Schedule") # Header English
    
    repayment_period_years_display = st.session_state.total_loan_term - st.session_state.grace_period
    st.subheader(f"{st.session_state.grace_period} YEARS GRACE, {repayment_period_years_display} YEARS PAYMENT, TOTAL {st.session_state.total_loan_term} YEARS LOAN TERM") # Subheader English
    st.write(f"**Loan Principal:** {format_number(st.session_state.principal_amount, currency_symbol=st.session_state.selected_currency_symbol)}") # Label English
    st.write(f"**Annual Interest Rate:** {format_number(st.session_state.interest_rate, is_percentage=True)}") # Label English and percentage format used
    st.markdown("---")

    # Display schedule as a DataFrame
    display_df = pd.DataFrame(st.session_state.schedule_data)
    
    # Rename DataFrame columns to English
    display_df = display_df.rename(columns={
        'YIL': 'YEAR',
        'ANAPARA ÖDEMESİ': 'PRINCIPAL PAYMENT',
        'FAİZ': 'INTEREST',
        'ANAPARA+FAİZ': 'P+I',
        'ÖDEME': 'PAYMENT',
        'KALAN BAKİYE': 'REMAINING BALANCE'
    })


    for col in display_df.columns:
        if col == 'YEAR':
            display_df[col] = display_df[col].apply(lambda x: format_number(x, is_year=True))
        else:
            display_df[col] = display_df[col].apply(lambda x: format_number(x, currency_symbol=st.session_state.selected_currency_symbol))

    st.dataframe(display_df.style.set_properties(**{'text-align': 'right'}), use_container_width=True)
    st.markdown("---")
    st.subheader(f"Total Amount Paid: {format_number(st.session_state.total_payment, currency_symbol=st.session_state.selected_currency_symbol)}") # Label English
    st.markdown("---")

    # --- Download Buttons ---
    st.subheader("Download Results") # Header English
    col_dl1, col_dl2 = st.columns(2)

    with col_dl1:
        excel_xlsx_data = create_excel_xlsx_report(
            st.session_state.schedule_data, 
            st.session_state.total_payment,
            st.session_state.principal_amount,
            st.session_state.interest_rate,
            st.session_state.grace_period,
            st.session_state.total_loan_term,
            st.session_state.selected_currency_symbol
        )
        st.download_button(
            label="Download as Excel (.xlsx)", # Button English
            data=excel_xlsx_data,
            file_name="loan_repayment_schedule.xlsx", # File name English
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="download_excel_xlsx_btn"
        )
        st.info("The Excel file will be formatted with currency symbols and automatic column widths.") # Info message English

    with col_dl2:
        word_docx_data = create_word_docx_report(
            st.session_state.schedule_data, 
            st.session_state.total_payment,
            st.session_state.principal_amount,
            st.session_state.interest_rate,
            st.session_state.grace_period,
            st.session_state.total_loan_term,
            st.session_state.selected_currency_symbol
        )
        st.download_button(
            label="Download as Word (.docx)", # Button English
            data=word_docx_data,
            file_name="loan_repayment_schedule.docx", # File name English
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_word_docx_btn"
        )